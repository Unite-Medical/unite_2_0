#!/usr/bin/env python3
"""Unite Medical — exhaustive page-by-page walkthrough & test book (docx).
Data-driven: every route is an entry; regenerate anytime the site changes."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x2A,0x4A); ACCENT=RGBColor(0x1A,0x6E,0x8E); GREEN=RGBColor(0x1E,0x7D,0x3A)
RED=RGBColor(0xB0,0x2A,0x2A); AMBER=RGBColor(0x9A,0x6A,0x00); GREY=RGBColor(0x55,0x55,0x55); WHITE=RGBColor(0xFF,0xFF,0xFF)

def shade(c,x):
    t=c._tc.get_or_add_tcPr(); s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),x); t.append(s)
def cell(c,t,bold=False,color=None,size=9,white=False,mono=False):
    c.text=''; p=c.paragraphs[0]; r=p.add_run(t); r.bold=bold; r.font.size=Pt(size)
    if mono: r.font.name='Courier New'
    if white: r.font.color.rgb=WHITE
    elif color: r.font.color.rgb=color
def hdr(t,hs,fill='0F2A4A'):
    for i,h in enumerate(t.rows[0].cells): cell(h,hs[i],bold=True,white=True,size=9); shade(h,fill)
def H1(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
def H2(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=ACCENT
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:spacing'); s.set(qn('w:before'),'260'); s.set(qn('w:after'),'80'); pr.append(s)
def H3(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=NAVY
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:spacing'); s.set(qn('w:before'),'160'); s.set(qn('w:after'),'40'); pr.append(s)
def body(d,t,italic=False,color=None,bold=False,size=10.5):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
def bullets(d,items):
    for it in items:
        p=d.add_paragraph(style='List Bullet'); r=p.add_run(it); r.font.size=Pt(10)
def mono_block(d,text,size=8):
    for line in text.strip('\n').split('\n'):
        p=d.add_paragraph(); pr=p._p.get_or_add_pPr()
        s=OxmlElement('w:spacing'); s.set(qn('w:after'),'0'); s.set(qn('w:line'),'220'); s.set(qn('w:lineRule'),'auto'); pr.append(s)
        r=p.add_run(line); r.font.name='Courier New'; r.font.size=Pt(size); r.font.color.rgb=NAVY
def pagebreak(d):
    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

STATUS_COLORS={'LIVE':GREEN,'DEMO':AMBER,'KEY':RED}

def page_entry(d, num, title, url, who, status, purpose, see, actions, tests, flow=None, note=None):
    # heading
    p=d.add_paragraph(); r=p.add_run(f"{num}  {title}"); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=NAVY
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:spacing'); s.set(qn('w:before'),'240'); s.set(qn('w:after'),'30'); pr.append(s)
    # meta line
    p=d.add_paragraph()
    r=p.add_run(url+'   '); r.font.name='Courier New'; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    r=p.add_run(f'[{status}]'); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=STATUS_COLORS.get(status.split()[0],GREY)
    r=p.add_run(f'   audience: {who}'); r.font.size=Pt(9.5); r.font.color.rgb=GREY
    # purpose (Damon-language)
    p=d.add_paragraph(); r=p.add_run('What this is (plain English): '); r.bold=True; r.font.size=Pt(10)
    r=p.add_run(purpose); r.font.size=Pt(10)
    if see:
        p=d.add_paragraph(); r=p.add_run('What you see: '); r.bold=True; r.font.size=Pt(10)
        r=p.add_run(see); r.font.size=Pt(10)
    if actions:
        p=d.add_paragraph(); r=p.add_run('What you can do here:'); r.bold=True; r.font.size=Pt(10)
        bullets(d, actions)
    if flow:
        p=d.add_paragraph(); r=p.add_run('The flow:'); r.bold=True; r.font.size=Pt(10)
        mono_block(d, flow)
    if note:
        p=d.add_paragraph(); r=p.add_run('⚠ Note: '); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=AMBER
        r=p.add_run(note); r.font.size=Pt(9.5); r.font.color.rgb=AMBER
    # test table
    t=d.add_table(rows=1,cols=2); t.style='Table Grid'; hdr(t,['✓','Test — pass means…'],'1A6E8E')
    for ts in tests:
        c=t.add_row().cells; cell(c[0],'[ ]',bold=True,size=10); cell(c[1],ts,size=9)

d=Document()
st=d.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
for sec in d.sections:
    sec.top_margin=Inches(0.6); sec.bottom_margin=Inches(0.6); sec.left_margin=Inches(0.7); sec.right_margin=Inches(0.7)

# ============================ COVER ============================
H1(d,'UNITE MEDICAL 2.0')
H1(d,'The Complete Walkthrough & Test Book')
body(d,'')
body(d,'Every page. Every flow. Every button worth pressing.',bold=True,size=13)
body(d,'')
body(d,'Prepared for the Damon + Alex in-person page-by-page review. This book covers all 111 routes '
       'on the site: what each page is for in plain English, what you can do on it, how the flows connect, '
       'and exactly what to test. Work front to back — public site first, then buying, then quoting, then '
       'portals, then the admin back office. Generated from the live codebase 2026-07-03 (commit 724a520).',size=11)
body(d,'')
H3(d,'Logins for the session')
t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,['Role','Email','Password'])
for a,b,c_ in [('Admin (Damon)','damon@unitemedical.net','admin'),
               ('Admin (Ops)','ops@unitemedical.net','admin'),
               ('Customer — surgical center','sarah@atlanta-surgical.com','demo'),
               ('Customer — pharmacy','kareem@holloway.com','demo')]:
    c=t.add_row().cells; cell(c[0],a,bold=True); cell(c[1],b,mono=True); cell(c[2],c_,mono=True)
body(d,'')
H3(d,'Status legend (top-right of every entry)')
t=d.add_table(rows=1,cols=2); t.style='Table Grid'; hdr(t,['Tag','Meaning'])
for a,b in [('LIVE','Fully real right now — real data, real APIs, no key needed'),
            ('DEMO','Works end-to-end on seeded demo data; flips to real automatically when its key/count lands'),
            ('KEY','Blocked on a specific credential (named in the entry) before it does anything real')]:
    c=t.add_row().cells; cell(c[0],a,bold=True,color=STATUS_COLORS.get(a)); cell(c[1],b)
body(d,'')
body(d,'House rule for the session: at every page ask the same three questions — Is it true? '
       'Would I trust it? Would I buy? Flag anything that fails any of the three.',italic=True,color=GREY)
pagebreak(d)

# ============================ MASTER MAP ============================
H2(d,'0. The whole site on one page')
body(d,'Three worlds, one platform. Customers see the left column. Logged-in buyers get the middle. '
       'You and the team run everything from the right.')
mono_block(d,'''
                 ┌───────────────────────  UNITEMEDICAL.NET  ───────────────────────┐
                 │                                                                   │
   PUBLIC SITE   │        CUSTOMER PORTALS            ADMIN BACK OFFICE (you)        │
   (anyone)      │        (logged-in buyers)          (damon@ / ops@)                │
─────────────────┼──────────────────────────────────────────────────────────────────┤
 Homepage        │  Dashboard                     Command center  /admin             │
 Catalog ────────┼─▶ Cart ─▶ Checkout ─▶ Order    Morning digest                     │
 Product pages   │  Track / Reorder               Orders ─▶ Fulfillment pipeline     │
 Source & Quote ─┼─▶ Self-serve quote portal      Quotes / Margin policy             │
 Diagnostics     │  Quote accept (/q/token)       Inventory + WMS (receive/lots/     │
 Robotics        │  Invoices / Quotes / Team      count/transfers) / Purchase orders │
 Services (4)    │  Rep portal (/rep)             Finance (AR) / Reps / CRM          │
 Segments (4)    │  Distributor portal (6 tabs)   Vendors / Discovery / Compliance   │
 Gov / About /   │                                Consignment / UDI / Surplus desk   │
 Compliance ...  │                                CMS / Settings / Integrations      │
 Free tools:     │                                Webhooks / AI dashboard            │
 Shortage list · │                                                                   │
 Supply risk ·   │                                                                   │
 Surplus market  │                                                                   │
─────────────────┴──────────────────────────────────────────────────────────────────┘
      47 public pages          10 account/portal surfaces         30 admin screens
''')
body(d,'')
body(d,'Suggested running order for your visit (roughly 2 half-days):',bold=True)
t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,['Block','Time','Chapters'])
for a,b,c_ in [('Day 1 AM — the storefront','2.5 h','Ch. 1–2: public pages + free tools'),
               ('Day 1 PM — spending money','2 h','Ch. 3–4: buy, quote, accept'),
               ('Day 2 AM — the portals','1.5 h','Ch. 5: account, rep, distributor'),
               ('Day 2 PM — running the company','2.5 h','Ch. 6: all 30 admin screens + Ch. 7 flows')]:
    c=t.add_row().cells; cell(c[0],a,bold=True); cell(c[1],b); cell(c[2],c_)
pagebreak(d)

# ============================ CH 1: PUBLIC MARKETING ============================
H2(d,'Chapter 1 — The public site (what a stranger sees)')
body(d,'Read these as a buyer who has never heard of Unite. No login. Every claim on these pages is a '
       'promise made in your name.')

E=[]
E.append(('1.1','Homepage','/','everyone','LIVE',
 'The front door. Its one job: in five seconds, a materials director understands what Unite is (veteran-owned '
 'medical supply + sourcing partner) and finds their door in (their segment, the catalog, or a quote).',
 'Full-bleed hero with live-inventory ticker, stats strip (GA warehouse · direct manufacturer relationships · 2pm same-day cutoff), '
 'partner logo marquee, six segment cards (surgery centers, pharmacies, EMS, health systems, distributors, government), '
 'shortage-list teaser, footer with credentials line.',
 ['Click each of the six segment cards — each must land on its matching page',
  'Click the primary CTA into the catalog','Click "Start a quote"','Use the top-nav mega menu — every link'],
 ['Hero renders in under ~2s and the first screen looks intentional on your phone too',
  'All six segment cards route correctly','Stats strip claims are ones you will stand behind',
  'Footer: FDA 3015727296 · CAGE 8MK70 · MSPV BPA 36C24123A0077 · DUNS 117553945 all correct',
  'No Nevada/Dallas references anywhere (one-warehouse footprint)'],None,
 'Partner logo bar may show placeholders — logos still being sourced; not a launch blocker.'))
E.append(('1.2','Catalog','/catalog','buyers','LIVE',
 'The store shelf: all 87 real products scraped from your live Shopify store, filterable by category '
 '(Orthotics, Diagnostics, PPE, Surgical, Supplements).',
 'Grid of product cards with photo, name, price; category filter tabs; search box; stock badges.',
 ['Filter by each of the 5 categories','Search a product you know by name','Click into 3+ products',
  'Check prices against what you actually charge'],
 ['Every category filter returns sensible products','Search finds products by partial name',
  'Prices, photos, pack sizes correct on spot-check of 5 SKUs you know cold',
  'Out-of-stock products show as such (demo stock counts until WMS opening count)']))
E.append(('1.3','Product page','/products/:id','buyers','LIVE',
 'One product, sold properly: photos, variants, price, description, and — the clever part — "Stocked '
 'equivalents": if this item is out or a buyer is comparing, we show substitutes we stock. That is the '
 'backorder-insurance pitch built into every page.',
 'Image gallery, variant picker, quantity, add-to-cart, tabbed description/specs, stocked-equivalents strip, reviews.',
 ['Add to cart with a variant + quantity','Open every image','Read the description for accuracy',
  'Click a stocked equivalent'],
 ['Variants switch price correctly','Add-to-cart respects available stock',
  'Stocked equivalents are genuinely comparable products (not random)',
  'Description is truthful — no claim you could not defend to FDA']))
E.append(('1.4','Source & Quote (start)','/quote','buyers','LIVE',
 'The front door for "we need something you do not stock." A buyer describes what they need — a brand, a spec, '
 'or a full custom/private-label run — and it lands on the sourcing desk as a lead. This is the top of the '
 'sourcing funnel, deliberately simple: no spreadsheets, no jargon.',
 'Three ways in: (a) name a product/brand/SKU + quantity, (b) describe a spec to build (materials, sizes, '
 'certifications, label preference: my label / Unite label), (c) link to /quote/new to upload a vendor sheet. '
 'Form promises pricing back inside one business day.',
 ['Submit a named-product request','Submit a custom-spec request','Follow the "Upload it here" link to /quote/new'],
 ['Both request types create a reference number and confirmation',
  'The one-business-day promise is one you will staff',
  'NO template-download button anywhere in this flow (removed per your review)',
  'Submission arrives in admin (check /admin/quotes or CRM leads)'],
 '''
  Buyer arrives with a need
      │
      ├── knows the product ──▶ name it + qty ──▶ reference # ──▶ sourcing desk (1 biz day)
      ├── needs it built ─────▶ spec form (materials/sizes/label) ──▶ same desk
      └── has a vendor sheet ─▶ /quote/new (Chapter 4) ──▶ instant engine quote
'''))
E.append(('1.5','Diagnostics line','/diagnostics','pharmacy/lab buyers','LIVE',
 'A dedicated storefront page for the diagnostics business — COVID/flu/strep/HIV rapid tests wholesale. '
 'Speaks to pharmacies and labs specifically; routes to the diagnostics desk for volume pricing.',
 'Category hero, featured test kits, volume-pricing pitch, contact-the-desk CTA.',
 ['Click through to featured diagnostic products','Use the "Contact our diagnostics desk" link'],
 ['Products shown are current, sellable test kits','Desk link opens contact with the right reason pre-filled']))
E.append(('1.6','Robotics program','/robotics','hospitals/ASCs','LIVE',
 'The flagship Restore Robotics partnership page: FDA 510(k)-cleared remanufactured da Vinci Xi/DV5 '
 'instruments at a fraction of new price, with Unite running logistics/consignment. You confirmed every claim '
 'with Restore — this page is where they live.',
 'Program explainer, how-it-works steps (sealed-container returns to Restore), FAQ (clearance, warranty, '
 'manufacturer of record), intake form for interested facilities.',
 ['Read every Restore claim once more','Submit the interest form','Check the FAQ answers'],
 ['"Industry\'s only FDA 510(k) clearance for remanufacturing da Vinci Xi and DV5" — verified with Restore ✓',
  'Warranty/manufacturer-of-record language matches the actual agreement',
  'Form submission promises follow-up within one business day — staffed?']))
E.append(('1.7','Services hub','/services','all B2B','LIVE',
 'The menu of what Unite does beyond the shelf: Distribution & Fulfillment, PDAC Consulting, Quoting & '
 'Sourcing, Distributor Program (plus robotics & diagnostics tiles).',
 'Card per service with a one-line pitch, each linking to its detail page.',
 ['Click every service card'],
 ['The services listed are the services you actually sell — nothing aspirational',
  'Each card routes to the right detail page']))
E.append(('1.8','Service · Distribution','/services/distribution','all B2B','LIVE',
 'The logistics pitch: one Georgia warehouse, same-day ship before 2pm EST, all 50 states, picked within 2 '
 'hours, tracking to your portal.',None,
 ['Read the operational promises as commitments'],
 ['2pm cutoff / same-day / 2-hour pick promises are real warehouse practice','No multi-warehouse claims']))
E.append(('1.9','Service · PDAC Consulting','/services/pdac','DME suppliers','LIVE',
 'YOUR page. #1 on Google for "PDAC consulting". Sells the coding/verification expertise: HCPCS coding for '
 'DME and orthotics, PDAC verification guidance. Every word carries your reputation.',
 'Positioning copy, what-you-get, engagement CTA.',
 ['Read every word aloud','Submit the consult request'],
 ['Nothing overpromises (no "guaranteed approval" language)','Request routes to you/CRM',
  'The page still ranks: after launch, watch Search Console weekly']))
E.append(('1.10','Service · Distributor Program','/services/distributors','regional distributors','LIVE',
 'Recruiting page for regional distributors who resell Unite lines: sourcing + fulfillment behind their brand.',
 None,['Read the terms as offers you honor','Apply as a fake distributor'],
 ['Program terms match how you actually run distributor deals (see also the distributor portal, Ch. 5)']))
E.append(('1.11','Service · Private Label','/services/private-label','brand owners','LIVE',
 'Sell the factory network: your product, your label, Unite runs spec→manufacture→landed delivery.',None,
 ['Submit the private-label inquiry'],
 ['Capability described = capability you have today','Inquiry lands on the sourcing desk']))
E.append(('1.12','Segments (4 pages)','/segments/asc · /pharmacy · /ems · /distributors','per segment','LIVE',
 'Four landing pages that greet each buyer type in their own language: surgery centers (case-cart staples, '
 'bracing, recovery DME), independent pharmacies (wholesale + diagnostics), EMS & fire (first-responder kit), '
 'regional distributors (sourcing muscle).',
 'Each: segment hero, curated product strips, segment-specific proof points, CTA into catalog/quote.',
 ['Read each page as that buyer','Click each CTA'],
 ['Each page uses that audience\'s vocabulary correctly (you know these buyers — does it ring true?)',
  'Product strips show segment-appropriate items','ASC page: "Net-30 on approved credit" phrasing is the conditional one ✓']))
E.append(('1.13','Government & VA','/government','contracting officers','LIVE',
 'The federal storefront: who Unite is to a contracting officer. Medava-branded SKUs on the MSPV BPA via an '
 'authorized SDVOSB distributor who holds the contract; Unite itself is veteran-owned, SAM-registered, CAGE 8MK70, ready to bid.',
 'Credential tiles (BPA · CAGE · Berry · TAA), contract-vehicle explainers, set-asides paragraph (your new wording), gov sales desk CTA.',
 ['Read the BPA attribution line — your exact requested wording','Check every credential value'],
 ['BPA line reads: Medava SKUs · via authorized SDVOSB distributor (contract holder) ✓',
  '"Ready to bid RFQs, IFBs, RFPs" — true and staffed','Gov desk contact routes correctly']))
E.append(('1.14','Procurement / diversity','/procurement','supplier-diversity buyers','LIVE',
 'For hospital supplier-diversity programs: Unite as the veteran-owned partner (DD214/ID.me verified), plus '
 'the diversity classifications your partner network holds.',None,
 ['Read the veteran positioning','Check DD214/ID.me framing'],
 ['No VOSB/SDVOSB self-claim anywhere — only partner attributions (verifier enforces this)']))
E.append(('1.15','Portfolio / outcomes','/portfolio','prospects','LIVE',
 'Proof page: seven numbered outcomes (including Medava SKUs on the national MSPV contract via the SDVOSB '
 'distributor) plus anonymized client-type cards.',None,
 ['Read all seven outcome blocks'],
 ['Every number/outcome is defensible','Anonymized cards stay anonymized (no accidental client names)']))
E.append(('1.16','Case study · TJS','/case-studies/tjs','physician groups','LIVE',
 'The Total Joint Specialists story: their branded patient recovery store, run on Unite rails. Named with '
 'permission; written to make the next physician group want one.',None,
 ['Read the numbers in the story','Confirm TJS is still happy being named'],
 ['Story numbers are real','The CTA for "want one of these" routes to contact/private-label']))
E.append(('1.17','About','/about','everyone','LIVE',
 'Your company, your letter. Est. 2019, Lithia Springs. The founder letter (2016 line removed, pandemic claim '
 'softened to "national scale"), leadership = you, the Unite family (Unite Pharma, Clyne Health).',
 'Hero + your letter + leadership card + family-of-companies strip + credentials grid.',
 ['Final read of your letter','Check the leadership section shows just you','Verify family-company descriptions'],
 ['Letter reads in your voice; nothing you would not say to a customer\'s face',
  'No Jackie card / no physician tagline (removed per your sign-off) ✓',
  'Founder photo placeholder — real warehouse shot still needed (known gap)']))
E.append(('1.18','Compliance','/compliance','quality/regulatory buyers','LIVE',
 'The trust page: FDA registration, BPA (new wording), CAGE, DUNS, TAA, Berry, PDAC tiles; quality-management '
 'statement (ISO 13485 "pursuing" — honest); recall/MDR posture; downloadable-document library with per-doc REQUEST buttons.',
 None,
 ['Check all 8 credential tiles','Click 2–3 document REQUEST buttons','Read the recall SLA'],
 ['Credential values all correct (they mirror the footer + PDFs)',
  'Document requests send to support@ and create a trackable request',
  '"Lot-level traceability… one business day" recall SLA — the WMS lot system backs this once live']))
E.append(('1.19','Locations','/locations','buyers','LIVE',
 'One dot done proudly: the Lithia Springs GA warehouse shipping all 50 states + territories, same-day before '
 '2pm. Spoke map, address, carrier list.',None,
 ['Look at the map','Check the address block'],
 ['ONE warehouse only — no Nevada remnants','1487 Trae Lane, Lithia Springs, GA 30122 correct']))
E.append(('1.20','Resources + HCPCS coding','/resources · /resources/coding','billers/buyers','LIVE',
 'Reference material that makes Unite look like the adults in the room: HCPCS Level II code reference for '
 'medical supplies, plus per-SKU coding reference. Feeds the PDAC positioning.',None,
 ['Look up 2 codes you know','Spot-check a SKU\'s suggested code'],
 ['Coding guidance is accurate — this page must never embarrass the PDAC brand']))
E.append(('1.21','Blog / field notes','/blog (+ /blog/:slug)','SEO/prospects','DEMO',
 'The content engine: market takes, compliance walkthroughs, ops notes. Renders clean but empty until Jill\'s '
 'articles land.',None,['Open the empty state'],
 ['Empty state looks intentional, not broken','Category filters render']))
E.append(('1.22','Careers','/careers','candidates','LIVE',
 'Simple hiring page.',None,['Skim listings'],['Nothing stale; contact route works']))
E.append(('1.23','Contact','/contact','everyone','KEY (Resend for real email)',
 'Call-us-we-answer page: 833.868.6483 (accounting ext. 3), support@ and accounting@ addresses, reason-coded '
 'message form promising a reply inside one business day.',None,
 ['Submit a real test message','Call the number once for theater'],
 ['Form creates a lead ID and confirmation','Once Resend key lands: the message actually arrives at support@',
  'No info@ anywhere ✓']))
E.append(('1.24','Support','/support','customers','LIVE',
 '"Answers, not tickets": FAQ-driven self-help with escalation to contact.',None,
 ['Read the FAQs','Try the escalation path'],
 ['FAQ answers match current policy (esp. the net-30-with-credit-approval one ✓)']))
E.append(('1.25','Legal (4 pages)','/privacy · /terms · /returns · /shipping','everyone','LIVE',
 'Privacy, Terms (Net-30/60 at our discretion with credit review), Returns (authorized returns to Lithia '
 'Springs), Shipping (same-day before 2pm EST).',None,
 ['Skim all four; read Returns fully — it is the page customers fight over'],
 ['Returns policy matches how you actually handle returns','Terms payment language matches D-03 (card-first, terms via credit review)']))
E.append(('1.26','Login / Register','/login · /register','customers','LIVE',
 'The account doors. Register: 2-minute B2B application (org, segment, spend) → auto-approval scoring → '
 'approved accounts get wholesale pricing, card/ACH from day one, flexible terms with approved credit. Login: '
 'standard sign-in ("saved lists, order history, dedicated rep").',None,
 ['Register a fake company end-to-end','Log out / log back in'],
 ['Register makes NO net-30-at-signup promise ✓ (your fix)','Auto-approval decides sensibly (real org names approve)',
  'Welcome email queues (sends for real once Resend lands)']))
for e in E: page_entry(d,*e)
pagebreak(d)

# ============================ CH 2: FREE TOOLS ============================
H2(d,'Chapter 2 — The free tools (lead magnets that work for you)')
body(d,'These three pages give strangers real value with no login — and every use hands Unite demand data '
       'and a warm lead. They are the growth engine; treat bugs here as revenue bugs.')
E=[]
E.append(('2.1','Shortage list matcher','/shortage-list','any buyer with a backorder problem','LIVE',
 'A buyer pastes (or uploads) their backorder list — the stuff their current supplier cannot deliver — and we '
 'instantly show which lines Unite can cover from stock, with suggested equivalents. Unmatched lines become '
 'sourcing-desk requests. Their paste is the payment: we learn exactly what the market is short on.',
 'Paste box / .csv-.txt upload, live line-by-line match results with stocked alternates, "send to sourcing desk" for misses.',
 ['Paste 5 real product names from memory','Upload a small CSV','Send an unmatched line to the desk'],
 ['Matches are genuinely comparable (a wrong match here embarrasses the whole pitch)',
  'Unmatched lines create shortage_requests + a CRM lead','Works on a phone — buyers do this from the floor'],
 '''
  Buyer's backorder list ──▶ paste/upload ──▶ instant match vs. 87-SKU catalog
        │                                         │
        │                              in stock ──▶ "we cover these today" + cart links
        │                              no match ──▶ sourcing desk request + CRM lead
        └────────────── every line = demand data we keep ──▶ future PRD-30 pricing brain
'''))
E.append(('2.2','Supply risk monitor','/supply-risk','proactive buyers','LIVE',
 'A live FDA recall/enforcement feed (openFDA — free, real, already on) mapped against the categories Unite '
 'stocks. When a recall hits a buyer\'s supplier, this page shows whether Unite covers the gap. It is the '
 '"ambulance-chaser with a heart" play: be there at the exact moment supply breaks.',
 'Recall cards (firm, reason, date) tagged with "we stock alternates" where true; deep links into shortage matcher.',
 ['Scan the current recalls','Click through a recall into the matcher'],
 ['Feed loads real, current FDA data','Category mapping is sensible','This is a page you would happily send a prospect']))
E.append(('2.3','Surplus intake + marketplace','/surplus · /surplus/market','sellers & bargain buyers','LIVE',
 'Two-sided: (a) organizations with excess/near-expiry inventory list it (manual or CSV upload) with target '
 'prices — AI normalizes the lines; the desk reviews, values, and makes offers. (b) Accepted lots publish to a '
 'public brokered marketplace where buyers bid.',
 'Intake: line-item form + CSV upload + how-it-works steps. Market: browsing of live lots with offer buttons.',
 ['List 2 fake surplus lines','Upload the surplus CSV','Browse the market and place a fake offer'],
 ['Intake normalizes messy product names sensibly (AI is live)','Offers flow to /admin/surplus for review',
  'Marketplace lots look presentable to a stranger'],
 '''
  Seller lists surplus ──▶ AI normalizes ──▶ /admin/surplus desk review ──▶ offer email
                                                        │ accepted
                                                        ▼
                                        lot publishes to /surplus/market ──▶ buyer offers
                                                        │ accept
                                                        ▼
                                     connection released · fee invoice · both sides emailed
'''))
for e in E: page_entry(d,*e)
pagebreak(d)

# ============================ CH 3: BUYING ============================
H2(d,'Chapter 3 — Spending money (the order lifecycle)')
body(d,'The single most important chain on the site. Do it twice: once as a fresh registrant, once as Sarah '
       '(demo customer). If anything here stutters, launch waits.')
mono_block(d,'''
 BROWSE ─▶ CART ─▶ CHECKOUT ─▶ CONFIRMED ─▶ (admin fulfills) ─▶ SHIPPED ─▶ TRACK ─▶ REORDER
   │        │         │            │                                │          │
 catalog  tier     stock-gate   order #        pipeline:         tracking   one-click
 search   pricing  + payment    + email        reserve→invoice→  number     repeat with
 filters  applied  (Stripe)     confirm        pack→ship→notify  to buyer   price check
''')
E=[]
E.append(('3.1','Cart','/cart','logged-in buyers','LIVE',
 'The basket. Tier pricing (A/B/C/distributor/gov) is applied per the account — two customers see different '
 'prices for the same item, on purpose. Out-of-stock items cannot be added at all.',
 'Line items with qty steppers, per-line and total pricing at the account tier, proceed-to-checkout.',
 ['Add 3 items as Sarah','Change quantities','Try to exceed available stock','Note the tier pricing'],
 ['Totals recalculate instantly and correctly','Stock gate refuses over-available quantities',
  'Tier discount matches the account type (Sarah = surgical center tier)']))
E.append(('3.2','Checkout','/checkout','buyers','KEY (Stripe for real payment)',
 'Payment + shipping capture. New accounts pay by card/ACH (your D-03 rule); accounts with approved credit see '
 'their terms. Live-stock re-check at submit so a sold-out item cannot slip through.',
 'Address, shipping method, payment method (card/ACH; terms only if approved), order review, place-order.',
 ['Complete an order end-to-end','Watch the confirmation page + email','Try checkout with an emptied cart (should refuse)'],
 ['Order lands with a real order number','Confirmation email queues (sends when Resend lands)',
  'Payment methods offered match the account\'s approved terms','Stock re-validated at submit']))
E.append(('3.3','Order confirmed + tracking','/orders/:id/confirmed · /orders/:id/track','buyers','DEMO',
 'The receipt page and the customer-facing status page (tracking number appears once ShipStation is keyed; '
 'demo status until then).',None,
 ['Open both pages for your test order'],
 ['Confirmation shows the right items/totals/rep contact','Track page presents status honestly (no fake tracking numbers)']))
E.append(('3.4','Reorder','/account/order · /account/reorder','repeat buyers','LIVE',
 'The repeat-purchase engine: pick a past order (or a saved list), preview the lines with current prices and '
 'stock, one click back into the cart. This is where B2B revenue compounds.',
 'Order history with per-order Reorder buttons, reorder preview modal (lines, current price, availability), add-all-to-cart.',
 ['Reorder your Chapter-3 test order','Watch the preview flag any price/stock changes'],
 ['Reorder preview matches the original order\'s lines','Price changes since last order are visible, not hidden',
  'Out-of-stock lines are flagged rather than silently dropped']))
for e in E: page_entry(d,*e)
pagebreak(d)

# ============================ CH 4: QUOTING ============================
H2(d,'Chapter 4 — Quoting & sourcing (the engine room)')
body(d,'Unite\'s real differentiator. Three distinct quote paths — know which is which, because a customer '
       'will only ever see two of them.')
mono_block(d,'''
 THREE QUOTE PATHS
 1. /quote        "I need X"          simple request form  → sourcing desk (customer-facing)
 2. /portal/quote "price my list"     self-serve catalog quoting at MY tier (customer-facing)
 3. /quote/new    vendor sheet in,    the landed-cost ENGINE: parse→translate→classify→
    (internal/rep) customer quote out  duty→freight→margin → branded PDF (internal + reps)

 All three converge:  QUOTE ──▶ /q/:token accept link ──▶ signed acceptance ──▶ ORDER
''')
E=[]
E.append(('4.1','Self-serve quote portal','/portal/quote','approved customers','LIVE',
 'Approved buyers build their own quote from the stocked catalog — priced at their tier automatically, '
 'including any per-SKU contract prices — generate it, and accept it online. No rep in the loop unless they '
 'want one. Gated to A-tier + approved distributors (D-09).',
 'Catalog search, add-SKU-with-qty lines, live tier pricing, generate-quote button, link to the accept page. '
 'Un-stocked needs route to the sourcing desk.',
 ['Build a 4-line quote as Sarah','Generate it','Follow the accept link'],
 ['Prices match the account tier / contract overrides','Generated quote appears in /account/quotes and /admin/quotes',
  'Un-stocked request becomes a sourcing lead']))
E.append(('4.2','Vendor-sheet engine','/quote/new','you + reps (internal)','LIVE',
 'The crown jewel. Upload any vendor price sheet — Excel or CSV, any language (Chinese/Korean/Vietnamese '
 'headers handled), any column order — and the engine: maps columns (with a confirmation step), translates, '
 'auto-classifies FDA product codes, pulls real USITC duty rates, compares LCL/FCL/air freight, applies '
 'tier margin with a floor, and emits an all-in landed-cost customer quote as a branded PDF. Multi-vendor '
 'compare picks the best FOB per product across sheets.',
 'Upload/paste (STEP 1 — the template download is gone per your review), column-mapping confirmation with '
 'confidence badges (exact/fuzzy/AI), parsed-line preview with translation, tier + freight preference pickers, '
 'run-engine progress, quote output with internal vs. customer view, Flexport classification CSV export.',
 ['Upload the sample (Load sample) and run the engine','Paste a messy CSV and watch column mapping ask for confirmation',
  'Run a foreign-language sheet if you have one','Compare two vendor sheets','Open both print views'],
 ['Engine completes on the sample in seconds','Column mapper flags uncertain columns instead of guessing silently',
  'Duty rates are real (USITC live) — spot-check one HTS code you know',
  'INTERNAL view shows cost breakdown; CUSTOMER view shows sell prices only — never the margins',
  'No "download our template" step exists ✓'],
 '''
 vendor.xlsx/csv (any language)
      │ parse (multi-sheet, 10MB guard)
      ▼
 column mapping ── alias → fuzzy → AI ──▶ you confirm uncertain ones
      ▼
 translate lines (original kept) ─▶ FDA code auto-classify ─▶ USITC duty lookup
      ▼
 landed cost = FOB + duty + ocean/air + brokerage + drayage + receiving
      ▼
 tier margin (A/B/C/dist/gov, 10% floor) ─▶ QUOTE ─▶ branded PDF ─▶ /q/:token accept
'''))
E.append(('4.3','Quote engine demo','/quote/engine','internal demo','DEMO',
 'The sanitized demo of the pipeline (run-sample button, step cards, sample quote). Kept for showing the '
 'capability without exposing internals. Not linked prominently for customers.',None,
 ['Run the sample once'],
 ['No internal tool names or margins visible anywhere in the customer-facing copy']))
E.append(('4.4','Quote accept (token link)','/q/:token','the customer','LIVE',
 'The magic link a customer clicks from their emailed quote: review lines, accept online, and the acceptance '
 'idempotently becomes a real order. No login required — the token IS the authorization.',None,
 ['Open the accept link from your 4.1 quote','Accept it','Try the same link again (should not double-order)'],
 ['Acceptance creates exactly one order (idempotent)','Accepted quote shows as accepted everywhere (account + admin)',
  'Expired quotes refuse gracefully']))
E.append(('4.5','Quote & invoice print/PDF','/quotes/:id/print · /invoices/:id/print','customers','LIVE',
 'The paper that lands on desks: branded quote (customer view + internal view) and invoice, with the footer '
 'carrying FDA · CAGE · MSPV BPA (your corrected wording). Print-to-PDF today; zero-dependency PDF engine behind it.',
 None,
 ['Download/print your test quote both views','Print an invoice'],
 ['Footer credentials correct on every document','Internal view NEVER goes to a customer (check the toggle default)',
  'Layout survives printing (no cut-off tables)']))
for e in E: page_entry(d,*e)
pagebreak(d)

# ============================ CH 5: PORTALS ============================
H2(d,'Chapter 5 — Account & partner portals')
body(d,'What logged-in customers, reps, and distributor partners each get as "their" Unite.')
E=[]
E.append(('5.1','Customer dashboard','/dashboard','customers','LIVE',
 'The customer\'s home: recent orders, open quotes, account snapshot, quick paths to reorder/track/quote.',None,
 ['Log in as Sarah and look around'],
 ['Numbers match her seeded history','Every tile links where it says']))
E.append(('5.2','Account · quotes / invoices / settings','/account/quotes · /account/invoices · /account/settings','customers','LIVE',
 'Self-service history: every quote (with accept/view), every invoice (with print), profile/org settings.',None,
 ['Open the quote from Ch. 4','Print an invoice','Change a setting'],
 ['Quote statuses accurate','Invoice math matches the order','Settings persist']))
E.append(('5.3','Account · team','/account/team','customer org owners','LIVE',
 'A customer\'s org can invite teammates with roles — owner / buyer / viewer. Buyers can order; viewers can '
 'look. Invitations activate when the teammate signs up. Owner-only controls.',None,
 ['Invite a fake teammate as Sarah','Change their role','Remove them'],
 ['Role gates actually gate (viewer cannot check out)','Pending invite activates on signup']))
E.append(('5.4','Rep portal','/rep','1099 sales reps','LIVE',
 'A rep\'s whole business on one page: book of business, attributed revenue and commission by window '
 '(30d/90d/1y), open quotes, recent orders, payout history, their Calendly intro link. Commissions accrue '
 'from orders automatically; payouts move real money once Stripe Connect is on.',None,
 ['Open as admin and inspect a rep\'s view','Check the commission math against /admin/reps'],
 ['Attribution logic matches the agreed flat-% structure (D-07)','Payout history consistent with the admin ledger']))
E.append(('5.5','Distributor portal (6 tabs)','/distributor','distributor partners','LIVE',
 'The partner cockpit for consignment/3PL distributors: My Inventory (their consigned stock), Order, Upload '
 'PO (drop their purchase order file, we parse it and draft the order), Shipping (their branded/blind-ship '
 'identities), Settlement (what they owe / are owed), Documents.',
 'Six-tab layout; PO-upload maps their SKUs to ours and learns the mapping for next time.',
 ['Walk each tab as the MedOne demo org','Upload a fake PO file','Check the settlement math'],
 ['PO upload parses and drafts order lines with SKU mapping','Blind-ship identity shows their brand, our address logic',
  'Settlement figures reconcile against the seeded orders',
  'This flow matches how the Restore/consignment deals actually run — you are the judge'],
 '''
 distributor uploads PO ─▶ parse + SKU-map (learns) ─▶ draft order ─▶ Unite fulfills
        │                                                    │ blind-ship under their brand
        ▼                                                    ▼
   their inventory view                              settlement ledger (fees, payouts)
'''))
for e in E: page_entry(d,*e)
pagebreak(d)

# ============================ CH 6: ADMIN ============================
H2(d,'Chapter 6 — The admin back office (30 screens, your operating system)')
body(d,'Log in as damon@unitemedical.net. This is where the company runs. For each screen: what it is, why '
       'you care, what to click. The daily rhythm:')
mono_block(d,'''
 YOUR MORNING          THE MONEY LOOP                    THE SUPPLY LOOP
 /admin/digest ──▶ /admin/orders ──▶ /admin/fulfillment   /admin/replenishment ─▶ draft POs
   5 ranked          new orders        reserve→invoice     ▼
   bullets w/        item search       →pack→ship→notify  /admin/purchase-orders ─▶ vendor
   deep links                                              ▼ (Flexport clears)
        │                                                 /admin/inventory/receive ─▶ stock
        ▼                                                  ▼ lots + QBO landed-cost bill
 /admin/finance (AR aging, reminders, record payments)    reorder points recalc — loop closes
''')
A=[]
A.append(('6.1','Command center','/admin','DEMO',
 'The landing view: today\'s orders, AR position, inventory alerts, recent activity. Your at-a-glance "is the '
 'business okay" screen.',
 ['Scan every tile','Click through 3 tiles to their detail screens'],
 ['Numbers agree with the detail pages behind them']))
A.append(('6.2','Morning digest','/admin/digest','LIVE (AI)',
 'The CEO brief: every morning, the system reads live orders/AR/inventory/freight/CRM/compliance and ranks '
 'the 5 things that most deserve your attention, each with a deep link. Claude writes it (key is on).',
 ['Generate today\'s digest','Judge: are these the right 5 things?'],
 ['Bullets are specific and actionable, not generic','Deep links land on the right screens']))
A.append(('6.3','Analytics','/admin/analytics','DEMO',
 'Sales and traffic analytics dashboards.',['Skim the charts'],['Charts render; ranges behave']))
A.append(('6.4','Orders','/admin/orders','LIVE',
 'Every order, searchable by item (find "who bought X"), with status and detail drill-in. Includes the '
 'inbound-shipment demo trigger for the receiving chain.',
 ['Find your Chapter-3 test order','Search orders by an item name','Open an order\'s full detail'],
 ['Item-level search works','Order detail shows payment/fulfillment state truthfully']))
A.append(('6.5','Fulfillment pipeline','/admin/fulfillment','LIVE',
 'Zero-touch order orchestration: validate → reserve stock → payment → invoice → shipping label → packing '
 'slip → notify → delivered, with per-step retries, circuit breakers per integration, backorder capture and '
 'auto-fulfill on restock, and returns/RMAs (restock + credit memo + refund).',
 ['Run your test order through the pipeline','Simulate a failure and watch the retry/backorder path','Process a fake return'],
 ['Each step shows observable state','A failed step retries rather than losing the order','RMA restocks and credits correctly']))
A.append(('6.6','Inventory home','/admin/inventory','DEMO (real after opening count)',
 'Stock per SKU per warehouse: on-hand, reserved, available, reorder points, movement history. The WMS '
 'projection — every number here is the sum of ledger movements, never hand-edited.',
 ['Browse a few SKUs','Open a movement history','Check available = on-hand − reserved'],
 ['Numbers labeled honestly as demo until the physical opening count seeds the ledger']))
A.append(('6.7','WMS · Receive','/admin/inventory/receive','LIVE (flow)',
 'The receiving workstation — built phone-first for the warehouse floor: scan/enter SKU + lot + expiry + '
 'quantity against an open PO; discrepancies flagged; every receipt writes lots + ledger + lands the landed-'
 'cost bill in QBO (when keyed).',
 ['Receive a line against an open PO on your phone','Enter a lot + expiry','Force a quantity mismatch'],
 ['Usable one-handed on a phone','Lot + expiry captured','Mismatch flags for review, does not silently accept']))
A.append(('6.8','WMS · Lots','/admin/inventory/lots','LIVE (flow)',
 'The lot browser + the recall box: look up any lot number and get every affected customer in under a second. '
 'This backs the /compliance one-business-day promise with a sub-second query.',
 ['Look up a seeded lot','Check expiring-soon list'],
 ['Recall lookup returns customers instantly','FEFO expiring-soon list is sensible']))
A.append(('6.9','WMS · Count','/admin/inventory/count','LIVE (flow)',
 'Cycle counting: open a session, count bins, variances post as ledger corrections after review.',
 ['Run a mini count with a deliberate variance'],['Variance posts a correction movement; on-hand updates']))
A.append(('6.10','WMS · Transfers','/admin/inventory/transfers','LIVE (flow)',
 'Warehouse-to-warehouse moves with an in-transit state (single-warehouse today; ready for warehouse #2).',
 ['Open the screen'],['Renders; creates a draft transfer']))
A.append(('6.11','Purchase orders','/admin/purchase-orders (+ print)','LIVE',
 'PO lifecycle board: draft → approved → sent → partial → received → closed. Replenishment drafts these '
 'automatically; you approve and send (email to vendor with the branded PO PDF).',
 ['Open a drafted PO','Approve + send one','Print the PO PDF'],
 ['Lifecycle statuses move correctly','PO PDF footer/address correct','Send queues the vendor email']))
A.append(('6.12','Replenishment','/admin/replenishment','LIVE',
 'The reorder brain: trailing-90-day run rate per SKU (Prophet forecast per-SKU when the sidecar deploys) → '
 'reorder points → days-of-cover → one-click draft POs grouped by vendor, MOQ-respected.',
 ['Sort by urgency','Draft POs from the suggestions','Sanity-check one suggested quantity'],
 ['Stockout/reorder/watch statuses make sense','Draft PO math = suggested qty × cost, grouped by vendor']))
A.append(('6.13','Products + editor','/admin/products (+ new/edit)','LIVE',
 'Full catalog CRUD: prices, images, descriptions, variants, visibility.',
 ['Edit a price and watch it flow to the storefront','Create a throwaway product and hide it'],
 ['Edits appear on the public catalog immediately','Visibility gating works']))
A.append(('6.14','Product onboarding','/admin/products/onboard','LIVE',
 'Pre-FDA-listing gauntlet for new products: GTIN check-digit + GS1 validation (key), openFDA classification, '
 'USITC duty — all validated before you ever open the FDA portal.',
 ['Run a fake product through'],['Each validation reports pass/fail with reasons']))
A.append(('6.15','Quotes desk','/admin/quotes','LIVE',
 'Every quote in flight from all three paths, with status and totals.',
 ['Find your Ch. 4 quotes'],['Statuses match reality (accepted shows accepted)']))
A.append(('6.16','Margin policy','/admin/settings/margin','LIVE',
 'The tier-margin editor: A 30 / B 50 / C 60 / distributor 25 / gov 20 (your D-02 numbers) with the 10% floor. '
 'Changing a number here changes every future quote.',
 ['Read the current numbers','Change one, quote, change it back'],
 ['Matches agreed policy','Change flows through to a fresh quote immediately']))
A.append(('6.17','Customers & approval','/admin/customers','LIVE',
 'Account list + the approval queue for registrations that scored MANUAL_REVIEW.',
 ['Approve/decline your fake Chapter-1 registration'],['Decision emails queue; account state changes']))
A.append(('6.18','CRM','/admin/crm (+ /admin/crm/hubspot)','KEY (HubSpot)',
 'Contacts and deals. Every order upserts the contact and logs a closed-won deal automatically. The HubSpot '
 'screen manages the two-way sync once the token lands.',
 ['Check your test order created a contact + deal'],['Order → CRM automation fired','HubSpot screen shows sync status']))
A.append(('6.19','Reps & payouts','/admin/reps','KEY (Stripe Connect for real money)',
 'The 1099 roster: attributed revenue, accrued commission, emailed statements, Calendly links, and Pay-via-'
 'Stripe (Express accounts, real transfers once Connect is approved; simulated ledger rows until then).',
 ['Review a rep\'s attribution','Send a statement','Run a simulated payout'],
 ['Commission math = flat % as agreed','Payout writes a ledger row and emails the statement']))
A.append(('6.20','Rep authority','/admin/team','LIVE',
 'What each rep is allowed to do — permission grants per rep.',
 ['Open a rep\'s grants'],['Grants reflect the intended authority model']))
A.append(('6.21','Vendor approval','/admin/vendors','LIVE',
 'The vendor gate: scored approval using live openFDA data (recalls, 510(k) history), weighted components, '
 'Class III auto-reject (your D-06 call), first-time-country manual review.',
 ['Score a fake vendor','Read the component breakdown'],
 ['Decisions come with visible reasons','Class III → AUTO_REJECT fires']))
A.append(('6.22','Trade discovery','/admin/discovery','DEMO (deferred key, D-14)',
 'Find manufacturers (vendor side) and US importers (customer side) from trade data; push either into the '
 'vendor pipeline or CRM with AI outreach drafts. Demo data until a trade-data provider is funded.',
 ['Run a search; push one lead to CRM'],['Clearly understood as sample data; workflow itself functions']))
A.append(('6.23','Finance / CFO','/admin/finance','KEY (QBO + Stripe for real books)',
 'AR aging buckets, one-click payment recording, one-click dunning reminders. The daily money screen.',
 ['Record a payment against your test order','Send (queue) a reminder'],
 ['Aging math correct','Payment recording updates the order + ledger']))
A.append(('6.24','Compliance / recalls','/admin/compliance','LIVE',
 'Continuous recall sweep across every stocked manufacturer (openFDA), mapped to affected SKUs, with AI-'
 'drafted customer notices into the outbox — review-gated.',
 ['Run a sweep','Read a drafted notice'],
 ['Sweep completes against real FDA data','Notices are review-gated (nothing auto-sends)']))
A.append(('6.25','Consignment','/admin/consignment','LIVE',
 'The consignment ledger behind the distributor/Restore programs: stock placed at partner sites, usage, '
 'settlement triggers.',
 ['Reconcile one consignment against the distributor portal view'],['Both sides agree']))
A.append(('6.26','UDI & GUDID','/admin/udi','LIVE',
 'Device-identifier operations: GS1 prefix capacity, DI issuance from medical-flagged prefixes, the post-'
 'quote gate queue, GUDID intake templates (Class 1+2), label-spec generation.',
 ['Walk one record through the gate queue','Generate a label spec'],
 ['DI issuance respects prefix rules','Gate queue statuses move correctly']))
A.append(('6.27','Surplus desk','/admin/surplus','LIVE',
 'Review surplus submissions, run AI valuation, adjust per-line offers, send the offer email, publish '
 'accepted lots to the public market.',
 ['Value your Chapter-2 fake submission','Send the offer','Publish a lot'],
 ['Valuation is reviewable before sending','Published lot appears on /surplus/market']))
A.append(('6.28','CMS','/admin/cms','LIVE',
 'Site content management for editable copy blocks.',['Change one block and view it live'],['Edit round-trips']))
A.append(('6.29','Settings & Integrations','/admin/settings · /admin/integrations (+ /ai)','LIVE',
 'Company settings (address/phone/BPA/CAGE feed the PDFs) and THE go-live board: every external service with '
 'configured-or-not status and a Run-a-ping test button. The AI screen shows per-prompt cost/latency/errors.',
 ['Ping every service','Read the AI cost dashboard'],
 ['Green pings = really talking to the upstream','Settings values match the site footer + PDFs',
  'This board is your single source of truth for "what is live"']))
A.append(('6.30','Webhooks','/admin/webhooks','LIVE',
 'The event bus console: incoming provider events (Stripe/ShipStation/Flexport/Calendly/Fathom), dedupe, '
 'retries with backoff, dead-letter queue, operator replay.',
 ['Inspect an event','Replay one from dead-letter (if any)'],
 ['Events show verified signatures once providers are keyed','Replay is idempotent (no double effects)']))
for num,title,url,status,purpose,actions,tests in A:
    page_entry(d,num,title,url,'internal (admin)',status,purpose,None,actions,tests)
pagebreak(d)

# ============================ CH 7: CROSS-CUTTING FLOWS ============================
H2(d,'Chapter 7 — The five flows that ARE the business (test end-to-end)')
body(d,'Individual pages passing is necessary, not sufficient. These chains crossing page boundaries are what '
       'actually make money. Run each start-to-finish without touching code.')

H3(d,'Flow A — Stranger → paying customer (the growth loop)')
mono_block(d,'''
 Google "PDAC consulting" ─▶ /services/pdac ─▶ /register (card-first) ─▶ approval
     or /supply-risk recall ─▶ /shortage-list paste ─▶ matched items ─▶ cart
                                                                         │
                    /catalog browse ─▶ product ─▶ cart ─▶ checkout ─▶ ORDER ─▶ email
''')
t=d.add_table(rows=1,cols=2); t.style='Table Grid'; hdr(t,['✓','End-to-end assertion'],'1A6E8E')
for x in ['A stranger can go from any public entry (SEO page, free tool, catalog) to a paid order without human help',
          'Every hop preserves context (shortage match → cart carries the matched items)',
          'The new-account order is card/ACH only; no terms leak']:
    c=t.add_row().cells; cell(c[0],'[ ]',bold=True,size=10); cell(c[1],x,size=9)

H3(d,'Flow B — Quote-to-cash (all three quote doors)')
mono_block(d,'''
 /quote request ──▶ desk follow-up ──▶ engine quote ─┐
 /portal/quote self-serve ───────────────────────────┼─▶ PDF + email ─▶ /q/:token ─▶ ACCEPT
 /quote/new vendor sheet (internal) ─────────────────┘                     │ (idempotent)
                                                              ORDER ─▶ fulfillment ─▶ invoice
''')
t=d.add_table(rows=1,cols=2); t.style='Table Grid'; hdr(t,['✓','End-to-end assertion'],'1A6E8E')
for x in ['Each of the three doors produces a quote that reaches /admin/quotes',
          'Accept link converts to exactly one order and marks the quote accepted everywhere',
          'Customer PDF never contains cost/margin columns']:
    c=t.add_row().cells; cell(c[0],'[ ]',bold=True,size=10); cell(c[1],x,size=9)

H3(d,'Flow C — Purchase-to-stock (the supply loop)')
mono_block(d,'''
 /admin/replenishment low-stock ─▶ draft PO ─▶ approve/send ─▶ vendor ─▶ Flexport ship
                                                                            │ customs cleared
   reorder points recalc ◀── QBO landed-cost bill ◀── lots+ledger ◀── /receive scan
   (the loop closes: stock in, books posted, next reorder smarter)
''')
t=d.add_table(rows=1,cols=2); t.style='Table Grid'; hdr(t,['✓','End-to-end assertion'],'1A6E8E')
for x in ['Simulated inbound shipment increments stock, posts the (stub) bill, and recalcs reorder points in one chain',
          'Received lot is immediately findable in the recall box',
          'PO shows partial when short-received']:
    c=t.add_row().cells; cell(c[0],'[ ]',bold=True,size=10); cell(c[1],x,size=9)

H3(d,'Flow D — Surplus two-sided loop')
mono_block(d,'''
 seller intake ─▶ AI normalize ─▶ desk valuation ─▶ offer email ─▶ accept
      ─▶ publish to market ─▶ buyer offer ─▶ accept ─▶ connection release + fee invoice
''')
t=d.add_table(rows=1,cols=2); t.style='Table Grid'; hdr(t,['✓','End-to-end assertion'],'1A6E8E')
for x in ['A submission travels intake → offer → market → buyer-accept without dead ends',
          'Competing offers close when one is accepted']:
    c=t.add_row().cells; cell(c[0],'[ ]',bold=True,size=10); cell(c[1],x,size=9)

H3(d,'Flow E — Distributor consignment loop')
mono_block(d,'''
 partner PO upload ─▶ SKU map (learns) ─▶ draft order ─▶ blind-ship under partner brand
                                   ─▶ consignment ledger ─▶ settlement ─▶ documents
''')
t=d.add_table(rows=1,cols=2); t.style='Table Grid'; hdr(t,['✓','End-to-end assertion'],'1A6E8E')
for x in ['Uploaded PO becomes a correct draft order including learned SKU mappings',
          'Settlement figures reconcile between admin consignment and the partner portal']:
    c=t.add_row().cells; cell(c[0],'[ ]',bold=True,size=10); cell(c[1],x,size=9)
pagebreak(d)

# ============================ CH 8: REDIRECTS + MATRIX ============================
H2(d,'Chapter 8 — Legacy URLs (protecting the Google rankings)')
body(d,'Every old URL must 301 to its successor or the SEO equity — especially the PDAC #1 — bleeds. Paste '
       'each into the browser and confirm where it lands.')
t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,['✓','Old URL','Must land on'])
for old,new in [('/solutions','/services'),('/solutions/asc','/segments/asc'),('/solutions/pharmacy','/segments/pharmacy'),
    ('/solutions/government','/government'),('/solutions/distributors','/segments/distributors'),('/solutions/ems','/segments/ems'),
    ('/services/dealer','/services/distributors'),('/services/education','/blog'),('/about/veteran-owned','/procurement'),
    ('/segments/gov','/government'),('any unknown URL','/ (homepage)')]:
    c=t.add_row().cells; cell(c[0],'[ ]',bold=True,size=10); cell(c[1],old,mono=True,size=9); cell(c[2],new,mono=True,size=9)

H2(d,'Chapter 9 — Master sign-off matrix')
body(d,'One line per chapter. Initial each when its chapter is fully walked and every flagged item is either '
       'fixed or explicitly accepted.')
t=d.add_table(rows=1,cols=4); t.style='Table Grid'; hdr(t,['Chapter','Pages/flows','Damon initials','Alex initials'])
for a,b in [('1 · Public site','26 entries'),('2 · Free tools','3 tools'),('3 · Buying','4 stages'),
            ('4 · Quoting','5 surfaces'),('5 · Portals','5 portals'),('6 · Admin','30 screens'),
            ('7 · End-to-end flows','5 chains'),('8 · Redirects','11 URLs')]:
    c=t.add_row().cells; cell(c[0],a,bold=True); cell(c[1],b); cell(c[2],'________'); cell(c[3],'________')
body(d,'')
body(d,'Final go decision (after Chapters 1–8 + the two money keys are live):',bold=True)
body(d,'')
body(d,'Damon (CEO):  ______________________________    Date: ____________')
body(d,'Alex (CTO):   ______________________________    Date: ____________')

d.save('/Users/alex-s-nt-16/Projects/Unite/unite_2_0/docs/Unite_Medical_Complete_Walkthrough_Test_Book.docx')
print('WROTE Test Book')
