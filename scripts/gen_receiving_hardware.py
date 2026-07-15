#!/usr/bin/env python3
"""Damon-facing hardware buy list for scan-to-receive (PRD-33).

Not a spec dump — a procurement card set Damon can read and act on: what to buy,
why, the trade-offs, and a blank DECISION line per item. Writes
docs/Receiving_Hardware.docx. Regenerate when hardware picks change.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0F, 0x2A, 0x4A); ACCENT = RGBColor(0x1A, 0x6E, 0x8E)
GREEN = RGBColor(0x1E, 0x7D, 0x3A); RED = RGBColor(0xB0, 0x2A, 0x2A)
AMBER = RGBColor(0x9A, 0x6A, 0x00); GREY = RGBColor(0x55, 0x55, 0x55); WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexc); tcPr.append(sh)


def set_cell(cell, text, bold=False, color=None, size=9, white=False):
    cell.text = ''; p = cell.paragraphs[0]; r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = WHITE
    elif color: r.font.color.rgb = color


def header_row(t, headers, fill='0F2A4A'):
    for i, h in enumerate(t.rows[0].cells): set_cell(h, headers[i], bold=True, white=True, size=9); shade(h, fill)


def h1(d, t):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY


def h2(d, t):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(13.5); r.font.color.rgb = ACCENT
    pPr = p._p.get_or_add_pPr(); s = OxmlElement('w:spacing'); s.set(qn('w:before'), '220'); s.set(qn('w:after'), '60'); pPr.append(s)


def h3(d, t):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = NAVY


def body(d, t, italic=False, color=None, bold=False, size=10.5):
    p = d.add_paragraph(); r = p.add_run(t); r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color


def bullet(d, t, lead=None):
    p = d.add_paragraph(style='List Bullet')
    if lead: rr = p.add_run(lead); rr.bold = True
    p.add_run(t)


d = Document()
st = d.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)
for sec in d.sections:
    sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)

h1(d, "Unite Medical — Receiving Hardware (Damon)")
body(d, "You asked for scan-to-receive: a worker scans the UPC on an incoming carton and it registers as "
        "received. That software is built and tested. This is the short list of what to BUY so it works on "
        "the warehouse floor. The software is scanner-agnostic — you are not locked to one brand. Read each "
        "card, pick a tier, and write your call on the DECISION line.", italic=True, color=GREY)

# ---- The one thing that matters most ----
h2(d, "The one decision that matters: 1D laser vs 2D imager")
body(d, "Every scanner types the barcode into the computer and presses Enter for you — that part is universal "
        "(it's called a 'USB keyboard wedge', zero setup). The real fork is WHAT it can read:")
t = d.add_table(rows=1, cols=3); t.style = 'Table Grid'
header_row(t, ["Type", "Reads", "Gets you"])
for kind, reads, gets in [
    ("1D laser (cheaper)", "The UPC barcode only (the striped code)",
     "Product recognized + units received. NO lot/expiry capture."),
    ("2D imager (recommended tier)", "UPC + 2D QR/DataMatrix codes, incl. medical UDI labels",
     "Product + LOT NUMBER + EXPIRATION auto-filled from one scan. This is what recalls and FEFO need."),
]:
    c = t.add_row().cells
    set_cell(c[0], kind, bold=True, size=9, color=NAVY); set_cell(c[1], reads, size=9); set_cell(c[2], gets, size=9)
body(d, "Why 2D matters for a medical supplier: device and drug cartons carry a 2D UDI code that packs the "
        "GTIN + lot + expiration together. A 2D scanner reads all three in one beep, so lot and expiry get "
        "captured automatically — that is the backbone of recall traceability and first-expire-first-out "
        "picking. A 1D laser can't see those codes, so a worker would type lot/expiry by hand.", size=10)
p = d.add_paragraph(); r = p.add_run("DECISION — scanner tier:  [ ] 2D imager (recommended)   [ ] 1D laser (UPC only)   [ ] not sure, talk it through")
r.bold = True; r.font.size = Pt(10)

# ---- Buy list ----
h2(d, "Buy list")
t = d.add_table(rows=1, cols=5); t.style = 'Table Grid'
header_row(t, ["Item", "What / why", "Example model", "Ballpark each", "Qty"])
for item, why, model, price, qty in [
    ("Barcode scanner (2D, corded USB)",
     "The main tool. Corded = never needs charging, plug into the receiving PC/tablet. 2D so it reads UDI lot+expiry.",
     "Zebra DS2208 or Honeywell Voyager 1450g2 (2D)", "$100 – $190", "1 per receiving station"),
    ("Barcode scanner (2D, wireless) — optional upgrade",
     "Same as above but cordless, for scanning cartons on a pallet away from the desk. Only if the receiving area is spread out.",
     "Zebra DS2278 (Bluetooth cradle)", "$230 – $320", "0 – 1"),
    ("Label printer (direct thermal)",
     "Prints Unite's own barcode labels — for blind receipts, repacks, or any carton whose factory barcode is missing/damaged. Thermal = no ink, cheap rolls.",
     "Zebra ZD230 (203 dpi, USB)", "$180 – $260", "1"),
    ("Label roll stock (4x2 or 4x3 in, thermal)",
     "Consumable for the printer above. Buy a couple boxes to start.",
     "4x2 direct-thermal, 1 in core", "$10 – $18 / roll", "2 – 4 boxes"),
]:
    c = t.add_row().cells
    set_cell(c[0], item, bold=True, size=9, color=NAVY); set_cell(c[1], why, size=8.5)
    set_cell(c[2], model, size=8.5); set_cell(c[3], price, size=9); set_cell(c[4], qty, size=9)
body(d, "Total to get one receiving station fully running with lot/expiry capture: roughly $290 – $450 "
        "(one 2D scanner + one thermal label printer + starter labels). Wireless upgrade adds ~$150.",
     bold=True, size=10)

# ---- What you do NOT need ----
h2(d, "What you do NOT need to buy")
bullet(d, "No special 'inventory software license' — the scanning + receiving is already in the Unite admin "
          "console (Receiving screen). The scanner just types into it.")
bullet(d, "No proprietary/brand-locked scanner. Anything sold as 'USB HID' or 'keyboard wedge' works. Avoid "
          "scanners that ONLY work with a vendor's own cloud app.")
bullet(d, "No handheld 'mobile computer' (the $1,000+ Zebra TC-series guns) for day one. A $150 corded 2D "
          "scanner at a fixed receiving station does the same job. Revisit only if you go to roaming/pallet scanning.")

# ---- How it works on the floor ----
h2(d, "How it works once it's plugged in")
for n, step in enumerate([
    "Carton arrives. Worker opens the Receiving screen and picks the purchase order it's for (or 'blind receipt' if there's no PO).",
    "Worker scans the barcode on the carton. Beep. The product name and SKU pop up automatically — no typing, no lookup.",
    "If it's a 2D UDI code, the lot number and expiration date fill in by themselves. Worker enters the quantity.",
    "Worker hits 'Post receipt'. It's now recorded in the system as received: on-hand goes up, a lot is created, and a scan record is logged for audit.",
    "The same screen shows a live RECONCILE panel: ordered vs received for that PO, with anything short, over, or unexpected flagged. When it all matches, it says BALANCED.",
], 1):
    p = d.add_paragraph(style='List Number'); p.add_run(step)
body(d, "That last part is the 'reconcile' you asked for — you can look at any incoming order and instantly "
        "see whether what showed up matches what was ordered. The received count is pulled from the actual "
        "stock ledger, not a number someone typed, so it can't be fudged.", italic=True, color=GREY, size=10)

# ---- One thing to know before go-live ----
h2(d, "One thing to flag before we flip it on for real")
t = d.add_table(rows=1, cols=2); t.style = 'Table Grid'
header_row(t, ["The gap", "What it takes"], fill='9A6A00')
c = t.add_row().cells
set_cell(c[0], "For a scan to recognize a product, that product needs its real UPC/GTIN stored in our system. "
              "The demo has placeholder codes. The live catalog (in the database) does not have UPCs yet.", size=9)
set_cell(c[1], "One-time load of the real manufacturer UPCs onto the product list. Best source is the actual "
              "codes on the cartons (Medava + others). I won't touch the live database without your go-ahead. "
              "Two ways: scan one sample carton per product to capture the true code, or pull them from a "
              "supplier feed. Your call on which.", size=9)
p = d.add_paragraph(); r = p.add_run("DECISION — how we load real UPCs:  [ ] scan a sample carton per SKU   "
                                     "[ ] get a code list from suppliers   [ ] Alex proposes, I approve")
r.bold = True; r.font.size = Pt(10)

d.save("docs/Receiving_Hardware.docx")
print("wrote docs/Receiving_Hardware.docx")
