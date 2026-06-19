#!/usr/bin/env python3
"""Generate the two interview-style .docx: Decisions + Keys/Accounts."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x2A,0x4A); ACCENT=RGBColor(0x1A,0x6E,0x8E)
GREEN=RGBColor(0x1E,0x7D,0x3A); RED=RGBColor(0xB0,0x2A,0x2A)
AMBER=RGBColor(0x9A,0x6A,0x00); GREY=RGBColor(0x55,0x55,0x55); WHITE=RGBColor(0xFF,0xFF,0xFF)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)
def set_cell(cell,text,bold=False,color=None,size=9,white=False):
    cell.text=''; p=cell.paragraphs[0]; run=p.add_run(text); run.bold=bold; run.font.size=Pt(size)
    if white: run.font.color.rgb=WHITE
    elif color: run.font.color.rgb=color
def header_row(t,headers,fill='0F2A4A'):
    for i,h in enumerate(t.rows[0].cells): set_cell(h,headers[i],bold=True,white=True,size=9); shade(h,fill)
def style_base(d):
    s=d.styles['Normal']; s.font.name='Calibri'; s.font.size=Pt(10.5)
def margins(d):
    for sec in d.sections:
        sec.top_margin=Inches(0.7); sec.bottom_margin=Inches(0.7)
        sec.left_margin=Inches(0.8); sec.right_margin=Inches(0.8)
def h1(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY; return p
def h2(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(13.5); r.font.color.rgb=ACCENT
    pPr=p._p.get_or_add_pPr(); spc=OxmlElement('w:spacing'); spc.set(qn('w:before'),'220'); spc.set(qn('w:after'),'60'); pPr.append(spc); return p
def body(d,t,italic=False,size=10.5,color=None,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return p
def bullet(d,t,lead=None):
    p=d.add_paragraph(style='List Bullet')
    if lead: rr=p.add_run(lead); rr.bold=True
    p.add_run(t); return p
def numbered(d,t):
    p=d.add_paragraph(style='List Number'); p.add_run(t); return p
def rule(d):
    p=d.add_paragraph(); p.add_run('_'*64).font.color.rgb=RGBColor(0xCC,0xCC,0xCC)

# ===========================================================
# DOC 1 — DECISIONS INTERVIEW
# ===========================================================
def decision(d, did, title, context, options, rec, env_effect):
    # title bar
    p=d.add_paragraph()
    r=p.add_run(f"{did}  ·  {title}"); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=NAVY
    pPr=p._p.get_or_add_pPr(); spc=OxmlElement('w:spacing'); spc.set(qn('w:before'),'240'); spc.set(qn('w:after'),'40'); pPr.append(spc)
    body(d,"Context: ",bold=True,size=10); 
    # context as its own paragraph
    cp=d.add_paragraph(); cr=cp.add_run(context); cr.font.size=Pt(10)
    # options table
    t=d.add_table(rows=1,cols=2); t.style='Table Grid'
    header_row(t,["Option","What it means"])
    for opt,mean in options:
        c=t.add_row().cells; set_cell(c[0],opt,bold=True,size=9,color=NAVY); set_cell(c[1],mean,size=9)
    # recommendation
    rp=d.add_paragraph(); rr=rp.add_run("Recommendation: "); rr.bold=True; rr.font.color.rgb=GREEN
    rr2=rp.add_run(rec); rr2.font.size=Pt(10); rr2.font.color.rgb=GREEN
    # what it wires
    if env_effect:
        ep=d.add_paragraph(); er=ep.add_run("What I wire when you choose: "); er.bold=True; er.font.size=Pt(9.5); er.font.color.rgb=GREY
        er2=ep.add_run(env_effect); er2.font.size=Pt(9.5); er2.font.color.rgb=GREY
    # answer line
    ap=d.add_paragraph(); ar=ap.add_run("YOUR DECISION: ____________________________________________    Date: __________")
    ar.font.size=Pt(10); ar.bold=True
    rule(d)

d=Document(); style_base(d); margins(d)
h1(d,"Unite Medical 2.0 — Decisions Interview")
body(d,"Every business decision needed before launch. We go through these together: "
       "I read the context + options + my recommendation, you choose, I wire it. "
       "Most of these are free (no key, no cost) — just your call. Last updated 2026-06-19.",
     italic=True,color=GREY)
body(d,"There are 14 decisions in two parts: A) business rules that shape how the platform "
       "behaves, and B) build/scope calls. Filling the 'YOUR DECISION' line is all we need.",size=10)

h2(d,"Part A — Business Rules")

decision(d,"D-01","Sales tax tooling",
 "The platform needs to charge correct sales tax per state at checkout. Three ways to do it.",
 [("Stripe Tax","Auto-calculates per state/product, ties to the Stripe we're already setting up. ~0.5% per transaction."),
  ("QBO built-in","Uses QuickBooks' own sales-tax engine. No extra fee but less automatic at checkout."),
  ("Avalara","Enterprise-grade, most accurate for complex nexus, but ~$50-100+/mo and heavier setup.")],
 "Stripe Tax — least setup, rides on the Stripe integration, fine for a 2-state footprint (GA+NV).",
 "Flip the checkout tax provider to Stripe Tax and enable it in the Stripe dashboard config.")

decision(d,"D-02","Margin policy per customer tier",
 "What markup we apply to each customer type. These defaults are already in the code; confirm or change.",
 [("Keep defaults","A (hospital/gov/retail) 30% · B (mid ASC/dealer) 50% · C (small clinic/one-off) 60% · Distributor 25%."),
  ("Adjust","Give me new numbers for any tier, or split further by product category.")],
 "Keep defaults for launch; revisit per-category after we see real margin data.",
 "Set the tier multipliers in /admin/settings/margin (marginPolicy.js).")

decision(d,"D-03","New-account payment terms",
 "When a new B2B customer is auto-approved, how can they pay on day one?",
 [("Credit-card only","New accounts start CC-only. Net-30 terms require a separate credit application later. (Safer.)"),
  ("Offer net-30 at signup","Riskier — extends credit before we know the customer.")],
 "Credit-card only at signup; net-30 via separate credit application. Standard B2B practice.",
 "Set the account-approval AUTO_APPROVE path to provision CC-only checkout (accountApproval.js).")

decision(d,"D-04","Quote validity period",
 "How long a generated quote's price stays locked before it expires.",
 [("14 days","Default. Customer-friendly."),
  ("7 days","Matches Flexport freight rates, which often expire in 7 days — protects our margin."),
  ("Other","Your number.")],
 "7 days if freight is a big cost component (protects landed-cost margin); 14 if you want it customer-friendly. Your call on the tradeoff.",
 "Set the quote-expiry constant in the quoting engine.")

decision(d,"D-05","Customer-facing margin disclosure",
 "On a quote, does the customer see our markup/landed-cost breakdown, or only the final sell price?",
 [("Hide (recommended)","Customer sees sell price only. Reps see the full landed-cost + margin breakdown internally."),
  ("Show","Full transparency to the customer. Unusual in distribution.")],
 "Hide from customer, show to rep. This is the normal distribution model.",
 "Set the quote-template flag so the customer PDF omits the internal cost columns.")

decision(d,"D-06","Vendor scoring & Class III device policy",
 "The vendor-approval engine auto-scores new vendors against live FDA data. Two sub-decisions.",
 [("Class III = AUTO_REJECT (v1)","High-risk implantable/life-sustaining devices get auto-rejected for launch. Safest."),
  ("Class III = manual review","Route to a human instead of auto-reject."),
  ("High-watch countries","Right now any first-time country goes to manual review. Want a stricter always-reject list for specific countries?")],
 "Class III AUTO_REJECT for v1; keep first-time-country manual review (no hard country blocklist yet).",
 "Confirm the scoring weights + Class III branch in vendorScoring.js.")

decision(d,"D-07","Rep commission structure",
 "How 1099 reps earn on attributed sales.",
 [("Flat % per rep","Each rep has one commission rate applied to their attributed revenue. Simple."),
  ("Tiered","Rate varies by segment or product category. More complex to administer.")],
 "Flat % per rep for launch; layer tiers later if needed.",
 "Set commission mode in the reps module (reps.js).")

decision(d,"D-08","Account-rep reassignment policy",
 "When an account moves to a different rep, what happens to the old rep's open deals?",
 [("Stay with old rep","Open deals stay credited to the original rep; only new business goes to the new rep. (Default.)"),
  ("Transfer all","Everything moves to the new rep immediately.")],
 "Stay with old rep — fairer and avoids commission disputes.",
 "Set the reassignment rule in the rep attribution logic.")

decision(d,"D-09","Self-serve quote gating",
 "Which customers can build & price their own quote without a rep (via /portal/quote)?",
 [("A-tier + approved distributors","Only large/trusted accounts self-serve. (Default.)"),
  ("All approved accounts","Anyone approved can self-serve."),
  ("Reps only","No self-serve; every quote goes through a rep.")],
 "A-tier + approved distributors. Protects pricing while serving your best accounts.",
 "Set the self-serve eligibility gate (selfServeQuote.js).")

decision(d,"D-10","Surplus buying terms",
 "When we buy surplus inventory from customers, our default offer + how we collect it.",
 [("35% of retail, contracted LTL","Offer 35% of retail for new-in-box Class II; we arrange LTL freight pickup. (Default.)"),
  ("Adjust %","Different baseline offer."),
  ("Customer-ships","Customer pays to ship it to us instead of contracted LTL.")],
 "35% of retail + contracted LTL pickup for launch.",
 "Set surplus offer baseline + logistics default (marketplace.js / surplus flow).")

decision(d,"D-11","Data retention & call-recording consent",
 "Compliance posture for stored data and (if we keep Fathom) call recording.",
 [("Approve defaults","Call transcripts: 3yr (closed deals) / indefinite (open). AI inputs: 30 days. Audit log: indefinite."),
  ("Adjust","Give me different retention windows.")],
 "Approve defaults; if we keep Fathom (see D-13) I add a consent line to /privacy.",
 "Set retention windows + add the privacy-policy consent clause if Fathom stays.")

h2(d,"Part B — Build / Scope Calls")

decision(d,"D-12","Lot-level tracking at launch",
 "Recall lookups in under a second, backing the compliance-page SLA. The database schema is ready; "
 "it needs the warehouse pick/pack flow to write a tracking row per shipped item — which is gated on Cin7 anyway.",
 [("Fast-follow","Launch without it; turn it on right after Cin7 is live. (Recommended.)"),
  ("Launch requirement","Block launch until lot tracking writes real rows.")],
 "Fast-follow — it depends on Cin7 being live, so it can't precede inventory go-live regardless.",
 "Sequence the lot-tracking wiring right after the Cin7 (B-10) milestone.")

decision(d,"D-13","Keep Fathom (call intelligence) for v1?",
 "Fathom auto-summarizes sales calls into CRM tasks. It's ~$19-29/user/mo and needs an in-call consent line. "
 "Nice-to-have, not core to selling.",
 [("Keep","Worth it — auto-captures call notes for the sales team."),
  ("Defer","Drop for v1, add later. Removes a cost + a consent requirement.")],
 "Defer for v1 unless the sales team specifically wants it now. One less thing to wire and consent for.",
 "If deferred, I leave the Fathom webhook stubbed and skip the privacy consent line.")

decision(d,"D-14","ImportGenius (trade discovery) for v1?",
 "Finds manufacturers + US importers as leads via global trade data. It's the most expensive single tool at "
 "$899/user/mo. The discovery screen works on demo data without it.",
 [("Provision now","Pay for it; turn on real trade-data lead discovery at launch."),
  ("Defer","Launch with the demo discovery screen; add when ROI is proven.")],
 "Defer — $899/mo is a lot to commit pre-launch. Add once the sales motion is proven.",
 "If deferred, /admin/discovery stays on demo data; flip the key in later with zero code change.")

decision(d,"D-15","Quote page — ship sanitized vs full rebuild",
 "The current /quote page is the older demo flow but fully sanitized (no proprietary leakage). The CTO spec "
 "calls for a full rebuild per Unite_Quoting_Engine_Spec.md. Note: the NEWER /quote/new upload→landed-cost flow "
 "is already the real engine.",
 [("Ship sanitized + push customers to /quote/new","Launch now; the real engine already exists at /quote/new."),
  ("Full rebuild first","Hold launch to rebuild the legacy /quote page per spec.")],
 "Ship sanitized and route quoting through /quote/new (which is the real landed-cost engine). Rebuild the legacy page post-launch.",
 "Point the primary 'Get a Quote' CTAs at /quote/new; leave the sanitized /quote as a fallback.")

# Summary capture table
h2(d,"Decision summary (we fill this as we go)")
t=d.add_table(rows=1,cols=3); t.style='Table Grid'
header_row(t,["ID","Your choice","Date"])
for did in ["D-01","D-02","D-03","D-04","D-05","D-06","D-07","D-08","D-09","D-10","D-11","D-12","D-13","D-14","D-15"]:
    c=t.add_row().cells; set_cell(c[0],did,bold=True,size=9,color=NAVY); set_cell(c[1]," ",size=9); set_cell(c[2]," ",size=9)

d.save("/Users/alex-s-nt-16/Projects/Unite/unite_2_0/docs/Unite_Medical_Decisions.docx")
print("WROTE Decisions")

# ===========================================================
# DOC 2 — KEYS / ACCOUNTS INTERVIEW
# ===========================================================
def keycard(d, kid, service, what, plan, envvars, steps, paste, priority):
    p=d.add_paragraph()
    r=p.add_run(f"{kid}  ·  {service}"); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=NAVY
    pr=p.add_run(f"     [{priority}]"); pr.bold=True; pr.font.size=Pt(9.5)
    pr.font.color.rgb = RED if priority=="LAUNCH BLOCKER" else (AMBER if priority=="HIGH" else GREY)
    pPr=p._p.get_or_add_pPr(); spc=OxmlElement('w:spacing'); spc.set(qn('w:before'),'240'); spc.set(qn('w:after'),'40'); pPr.append(spc)
    body(d,what,size=10)
    # meta table
    t=d.add_table(rows=1,cols=2); t.style='Table Grid'
    header_row(t,["Plan / cost","Env var(s) to set"])
    c=t.add_row().cells; set_cell(c[0],plan,size=9); set_cell(c[1],envvars,size=9,color=NAVY,bold=True)
    # steps
    sp=d.add_paragraph(); sr=sp.add_run("How to get it:"); sr.bold=True; sr.font.size=Pt(10)
    for s in steps: numbered(d,s)
    # paste
    pp=d.add_paragraph(); ppr=pp.add_run("Where it goes: "); ppr.bold=True; ppr.font.size=Pt(9.5); ppr.font.color.rgb=GREY
    ppr2=pp.add_run(paste); ppr2.font.size=Pt(9.5); ppr2.font.color.rgb=GREY
    # status line
    ap=d.add_paragraph(); ar=ap.add_run("STATUS:  [ ] account created   [ ] key obtained   [ ] pasted to Vercel   [ ] pinged green")
    ar.font.size=Pt(9.5); ar.bold=True
    rule(d)

d=Document(); style_base(d); margins(d)
h1(d,"Unite Medical 2.0 — Keys, Accounts & Setup Guide")
body(d,"Every external account + API key needed to take the site from demo to live, with step-by-step "
       "instructions for obtaining each one. We go through these together: you create the account / pull "
       "the key, hand it to me, I wire it into the environment and verify the ping goes green. "
       "Last updated 2026-06-19.",italic=True,color=GREY)

h2(d,"How the wiring works (read once)")
bullet(d,"The browser never holds a secret. Keys live server-side in Vercel → Settings → Environment Variables.")
bullet(d,"Adding a key is the ONLY thing that flips a surface from demo to live — no code change.")
bullet(d,"After each key lands, verify at /admin/integrations → 'Run a ping'. Green = the real upstream answered.")
bullet(d,"Order below = launch criticality. The two true blockers are K-01 (Resend email) and K-02 (Stripe).")

h2(d,"Already done")
t=d.add_table(rows=1,cols=2); t.style='Table Grid'; header_row(t,["Service","Status"])
for s,st in [("Anthropic / Claude (AI features)","DONE — key set"),
             ("Neon Postgres (durable database)","DONE — DATABASE_URL + sync tokens set"),
             ("openFDA (recall data) + USITC (duty math)","DONE — free, no key needed")]:
    c=t.add_row().cells; set_cell(c[0],s,size=9); set_cell(c[1],st,size=9,color=GREEN,bold=True)

h2(d,"Part 1 — Launch Blockers (do these first)")

keycard(d,"K-01","Resend (transactional email)",
 "Powers EVERY outbound email: order confirmations, shipping notices, invoices, dunning/reminders, "
 "purchase orders, rep statements. Until this key is set, all email silently queues in an outbox and "
 "nothing reaches customers. This is the single most important key.",
 "Free for 100 emails/day; ~$20/mo production tier",
 "RESEND_API_KEY",
 ["Go to resend.com and create an account (use a company email).",
  "In the dashboard go to Domains → Add Domain → enter unitemedical.net.",
  "Resend shows SPF, DKIM, and DMARC DNS records. Add these in your domain's DNS (where unitemedical.net is registered). This proves you own the domain so mail isn't marked spam.",
  "Wait for Resend to verify the domain (usually minutes, sometimes up to an hour).",
  "Go to API Keys → Create API Key → copy the key (starts with 're_'). Hand it to Alex."],
 "Vercel env var RESEND_API_KEY. Verify: send a test from /admin or trigger an order confirmation.",
 "LAUNCH BLOCKER")

keycard(d,"K-02","Stripe (payments, invoicing, rep payouts)",
 "Real customer payments at checkout, invoicing/AR, and 1099 rep payouts via Connect. No key = no real "
 "orders can be paid. Connect approval has lead time, so start today.",
 "Pay-per-transaction (~2.9% + 30c); Connect + Billing included",
 "STRIPE_SECRET_KEY · STRIPE_WEBHOOK_SECRET",
 ["Log into your existing Stripe account at dashboard.stripe.com (or create one — you'll need the federal EIN).",
  "Toggle to Live mode (top right). Go to Developers → API keys → copy the Secret key (starts with 'sk_live_').",
  "Start Connect approval: Settings → Connect → enable Express. Stripe asks for business details (EIN, bank). This can take a day or two — start it now.",
  "Set up the webhook: Developers → Webhooks → Add endpoint → URL https://<your-vercel-domain>/api/hooks/stripe → select events (checkout, invoice, payment_intent). Copy the signing secret (starts with 'whsec_').",
  "Hand Alex the secret key + the webhook signing secret."],
 "Vercel env vars STRIPE_SECRET_KEY and STRIPE_WEBHOOK_SECRET. Verify: /admin/integrations ping + a test checkout.",
 "LAUNCH BLOCKER")

h2(d,"Part 2 — Core Operations (high priority)")

keycard(d,"K-03","Cin7 Core (inventory / WMS)",
 "The real stock numbers the whole site trusts — catalog availability, checkout stock-gating, replenishment. "
 "Until this is live, inventory runs on demo numbers.",
 "~$349-799/mo depending on tier; includes onboarding help",
 "CIN7_ACCOUNT_ID · CIN7_APPLICATION_KEY",
 ["Contact Cin7 sales (cin7.com) and pick a Core plan. Their onboarding team helps migrate inventory + POs from Shopify.",
  "Once provisioned, in Cin7 go to Integrations & API → API v1 → generate credentials.",
  "Copy the Account ID and Application Key. Hand both to Alex.",
  "Coordinate the data migration window with the warehouse lead + Cin7 onboarding."],
 "Vercel env vars CIN7_ACCOUNT_ID + CIN7_APPLICATION_KEY. Verify: /admin/inventory shows live stock.",
 "HIGH")

keycard(d,"K-04","ShipStation (labels & tracking)",
 "Creates shipping labels and pushes tracking numbers to customers. You likely already have this via the "
 "Shopify plugin — you need the DIRECT API key, not the plugin-mediated one.",
 "Existing subscription",
 "SHIPSTATION_API_KEY · SHIPSTATION_API_SECRET · SHIPSTATION_WEBHOOK_SECRET",
 ["Log into ShipStation → Settings (gear) → Account → API Settings.",
  "Click 'Generate API Keys' if none exist. Copy the API Key and API Secret.",
  "For tracking webhooks: Settings → Integrations → Webhooks → add https://<your-vercel-domain>/api/hooks/shipstation. Set/copy the webhook secret.",
  "Hand Alex the key, secret, and webhook secret."],
 "Vercel env vars SHIPSTATION_API_KEY + SHIPSTATION_API_SECRET + SHIPSTATION_WEBHOOK_SECRET.",
 "HIGH")

keycard(d,"K-05","HubSpot (CRM)",
 "Contacts, deals, and rep pipeline — the sales backbone. Orders auto-create contacts + closed-won deals once wired.",
 "Sales Hub Pro ~$90/seat/mo (or Enterprise $150)",
 "HUBSPOT_PRIVATE_APP_TOKEN",
 ["In HubSpot go to Settings → Integrations → Private Apps → Create a private app.",
  "Name it 'Unite Medical Site'. Under Scopes, grant CRM read/write (contacts, deals, companies, tickets).",
  "Create the app → copy the Access Token. Hand it to Alex.",
  "Confirm how many rep seats you're buying (affects cost)."],
 "Vercel env var HUBSPOT_PRIVATE_APP_TOKEN. Verify: place a test order → contact appears in HubSpot.",
 "HIGH")

keycard(d,"K-06","Calendly (scheduling)",
 "Rep booking links; bookings flow into the CRM. Google Calendar is NOT required — Calendly covers scheduling.",
 "Pro ~$12/user/mo (need the API-access tier)",
 "CALENDLY_API_KEY · CALENDLY_WEBHOOK_SECRET",
 ["Upgrade your Calendly to a plan with API access (Pro or higher).",
  "Go to Integrations → API & Webhooks → generate a Personal Access Token. Copy it.",
  "Create a webhook subscription pointing to https://<your-vercel-domain>/api/hooks/calendly; copy/set the signing key.",
  "Hand Alex the token + webhook secret."],
 "Vercel env vars CALENDLY_API_KEY + CALENDLY_WEBHOOK_SECRET.",
 "HIGH")

keycard(d,"K-07","QuickBooks Online (accounting)",
 "Orders and invoices flow into the books — no double entry. Needs a one-time OAuth consent (not just a key).",
 "QBO Essentials $35/mo or Plus $65/mo + Intuit developer app (free)",
 "QBO_CLIENT_ID · QBO_CLIENT_SECRET · QBO_REALM_ID · QBO_REFRESH_TOKEN · QBO_ENVIRONMENT",
 ["Have your CFO confirm the QBO Online subscription (and plan migration from Desktop if applicable).",
  "Go to developer.intuit.com → create an app → get production Client ID + Client Secret.",
  "In the app settings, add redirect URI: https://<your-vercel-domain>/api/auth/qbo/callback.",
  "Give Alex the Client ID + Secret. He sets them + QBO_ENVIRONMENT=production, then you visit https://<your-vercel-domain>/api/auth/qbo/connect once and approve — that mints the Realm ID + Refresh Token automatically."],
 "Vercel env vars QBO_CLIENT_ID + QBO_CLIENT_SECRET + QBO_ENVIRONMENT, then OAuth fills the rest.",
 "HIGH")

keycard(d,"K-08","Flexport (freight & customs)",
 "Inbound ocean freight tracking + landed-cost into the books. Needs your Flexport rep to enable API access.",
 "Account-required; CSM enables the Public API tier",
 "FLEXPORT_API_KEY · FLEXPORT_WEBHOOK_SECRET",
 ["Email your Flexport Client Success Manager and ask for Public API access.",
  "Once enabled, generate an API token in the Flexport dashboard (Settings → API).",
  "Register webhook https://<your-vercel-domain>/api/hooks/flexport; copy/set the secret.",
  "Hand Alex the token + webhook secret."],
 "Vercel env vars FLEXPORT_API_KEY + FLEXPORT_WEBHOOK_SECRET.",
 "MEDIUM")

keycard(d,"K-09","GS1 US (barcode validation)",
 "Validates GTINs/barcodes during product onboarding before FDA listing.",
 "~$500/yr+ (you may already have this if registered with GS1)",
 "GS1_API_KEY · GS1_ACCOUNT_ID",
 ["Check if Unite already has a GS1 US Data Hub account (gs1us.org). If so, log in.",
  "Go to the Data Hub API / developer section → generate API credentials.",
  "Copy the API Key + Account ID. Hand both to Alex.",
  "If not registered, start a GS1 US membership — note this also governs your real barcodes."],
 "Vercel env vars GS1_API_KEY + GS1_ACCOUNT_ID. Verify: /admin/products/onboard GTIN check.",
 "MEDIUM")

h2(d,"Part 3 — Optional / Decision-gated")
body(d,"These depend on decisions D-13 (Fathom) and D-14 (ImportGenius) in the Decisions doc. "
       "Only pursue the key if the decision is 'keep/provision'.",italic=True,color=GREY)

keycard(d,"K-10","Fathom (call intelligence) — only if D-13 = keep",
 "Auto-summarizes sales calls into CRM tasks. Skip entirely if D-13 = defer.",
 "$19-29/user/mo",
 "FATHOM_WEBHOOK_SECRET",
 ["If keeping: set up a Fathom team account (fathom.video).",
  "Configure a webhook to https://<your-vercel-domain>/api/hooks/fathom; copy/set the secret.",
  "Hand Alex the webhook secret. (Also triggers adding a consent line to /privacy.)"],
 "Vercel env var FATHOM_WEBHOOK_SECRET.",
 "OPTIONAL")

keycard(d,"K-11","ImportGenius (trade discovery) — only if D-14 = provision",
 "Finds manufacturers + US importers as leads. Skip if D-14 = defer; the discovery screen runs on demo data without it.",
 "$899/user/mo (one shared user for v1)",
 "IMPORTGENIUS_API_KEY",
 ["If provisioning: subscribe at importgenius.com (Enterprise tier for API).",
  "Request API access from their team; generate an API key.",
  "Hand Alex the key."],
 "Vercel env var IMPORTGENIUS_API_KEY. Verify: /admin/discovery returns live trade data.",
 "OPTIONAL")

keycard(d,"K-12","Forecasting sidecar — deploy (not a key)",
 "Smarter reorder points via the Prophet model. This is the only surface that needs a container DEPLOYED, "
 "not just a key. Run-rate stub is fine for launch (per decision D-12 context).",
 "Hosting ~$5-15/mo (Fly.io / Render / Railway)",
 "FORECAST_API_URL · FORECAST_API_TOKEN (optional)",
 ["Deploy the forecasting/ folder (Dockerfile included) to Fly.io, Render, or Railway.",
  "Point it at the same DATABASE_URL.",
  "Set FORECAST_API_URL (the deployed container's URL) + optional FORECAST_API_TOKEN in Vercel.",
  "Alex can run the deploy — flag it when you want forecasting live."],
 "Vercel env vars FORECAST_API_URL + FORECAST_API_TOKEN. Verify: /admin/replenishment shows PROPHET badge.",
 "LATER")

# Master checklist
h2(d,"Master procurement checklist")
t=d.add_table(rows=1,cols=4); t.style='Table Grid'
header_row(t,["Done?","Service","Env var(s)","Priority"])
rows=[("K-01","Resend","RESEND_API_KEY","LAUNCH BLOCKER"),
 ("K-02","Stripe","STRIPE_SECRET_KEY + webhook","LAUNCH BLOCKER"),
 ("K-03","Cin7 Core","CIN7_ACCOUNT_ID + APPLICATION_KEY","HIGH"),
 ("K-04","ShipStation","SHIPSTATION key+secret+webhook","HIGH"),
 ("K-05","HubSpot","HUBSPOT_PRIVATE_APP_TOKEN","HIGH"),
 ("K-06","Calendly","CALENDLY_API_KEY + webhook","HIGH"),
 ("K-07","QuickBooks","QBO_* (OAuth)","HIGH"),
 ("K-08","Flexport","FLEXPORT_API_KEY + webhook","MEDIUM"),
 ("K-09","GS1","GS1_API_KEY + ACCOUNT_ID","MEDIUM"),
 ("K-10","Fathom (if D-13=keep)","FATHOM_WEBHOOK_SECRET","OPTIONAL"),
 ("K-11","ImportGenius (if D-14)","IMPORTGENIUS_API_KEY","OPTIONAL"),
 ("K-12","Forecast (deploy)","FORECAST_API_URL","LATER")]
for kid,svc,ev,pri in rows:
    c=t.add_row().cells
    set_cell(c[0],"[ ]",size=9,bold=True)
    set_cell(c[1],f"{kid} {svc}",size=9)
    set_cell(c[2],ev,size=8,color=NAVY)
    col=RED if pri=="LAUNCH BLOCKER" else (AMBER if pri=="HIGH" else GREY)
    set_cell(c[3],pri,size=8,color=col,bold=True)

d.save("/Users/alex-s-nt-16/Projects/Unite/unite_2_0/docs/Unite_Medical_Keys_and_Accounts.docx")
print("WROTE Keys")
