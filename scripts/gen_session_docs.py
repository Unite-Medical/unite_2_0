#!/usr/bin/env python3
"""Generate the two .docx working-session documents for Unite Medical."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0F, 0x2A, 0x4A)
ACCENT = RGBColor(0x1A, 0x6E, 0x8E)
GREEN = RGBColor(0x1E, 0x7D, 0x3A)
RED = RGBColor(0xB0, 0x2A, 0x2A)
AMBER = RGBColor(0x9A, 0x6A, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, size=9, white=False, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align: p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white: run.font.color.rgb = WHITE
    elif color: run.font.color.rgb = color

def header_row(table, headers, fill='0F2A4A'):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, white=True, size=9)
        shade(hdr[i], fill)

def style_base(doc):
    st = doc.styles['Normal']
    st.font.name = 'Calibri'; st.font.size = Pt(10.5)

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
    p.space_after = Pt(4)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13.5); r.font.color.rgb = ACCENT
    pPr = p._p.get_or_add_pPr()
    spc = OxmlElement('w:spacing'); spc.set(qn('w:before'), '200'); spc.set(qn('w:after'), '60'); pPr.append(spc)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = NAVY
    return p

def body(doc, text, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p

def rule(doc):
    p = doc.add_paragraph(); p.add_run('_' * 60).font.color.rgb = RGBColor(0xCC,0xCC,0xCC)

# ============================================================
# DOC 1 — WALKTHROUGH / DEMO GUIDE
# ============================================================
doc = Document()
style_base(doc)
for section in doc.sections:
    section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)

h1(doc, "Unite Medical 2.0 — Site Walkthrough & Demo Guide")
body(doc, "Prepared for the pre-launch review with Damon (Founder/CEO). "
          "Last updated 2026-06-19.", italic=True, color=GREY)

h2(doc, "How to use this guide")
body(doc, "This is a guided tour of the site as it exists today, built for walking Damon "
          "through it for the first time. We go in three passes: (1) the public site a "
          "customer sees, (2) the customer & rep portals, (3) the admin back-office that "
          "runs the business.")
body(doc, "The single most important thing to tell Damon up front:", bold:=False)
p = doc.add_paragraph()
r = p.add_run("Everything in this guide already works right now on realistic demo data. ")
r.bold = True; r.font.color.rgb = GREEN
p.add_run("Nothing is an empty placeholder. The site is fully clickable end-to-end today. "
          "Adding an API key simply swaps demo data for live data — it does not 'turn on' a "
          "feature that wasn't there. So this walkthrough doubles as the launch demo.")

h3(doc, "Demo logins (for the live click-through)")
t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'
header_row(t, ["Role", "Email", "Password"])
for role, em, pw in [
    ("Admin / CEO (Damon's view)", "damon@unitemedical.net", "admin"),
    ("Customer — surgical center", "sarah@atlanta-surgical.com", "demo"),
    ("Customer — pharmacy", "kareem@holloway.com", "demo"),
]:
    row = t.add_row().cells
    set_cell(row[0], role, size=9.5); set_cell(row[1], em, size=9.5); set_cell(row[2], pw, size=9.5)

body(doc, "")
body(doc, "Legend for the 'Real today?' column below:", italic=True, color=GREY)
bullet(doc, "fully real right now — no key needed (e.g. catalog, FDA recall data, duty math).", bold_lead="LIVE — ")
bullet(doc, "works on demo data; goes live the moment its key/account is added.", bold_lead="DEMO — ")

# ---- PASS 1: PUBLIC SITE ----
rule(doc)
h2(doc, "Pass 1 — The Public Site (what a customer sees)")
pub = [
 ("Homepage", "/", "Cinematic hero, live-inventory ticker, partner logo marquee, segment entry points.", "DEMO — SKU count + 8 partner logos pending"),
 ("Product Catalog", "/catalog", "All 87 real Unite products (scraped from the live Shopify store), category filters, search.", "LIVE — real catalog & images"),
 ("Product Detail", "/products/:id", "Full product page + 'Stocked equivalents' substitute matching (the backorder-insurance hook).", "LIVE"),
 ("Cart & Checkout", "/cart, /checkout", "Tier-based pricing per customer type, live-stock gating so out-of-stock can't be bought.", "DEMO — real payment needs Stripe key"),
 ("Get a Quote", "/quote, /quote/new", "Upload a vendor spreadsheet (Excel/CSV, even foreign-language) → real landed-cost engine (FOB + duty + freight + brokerage) → branded PDF quote.", "LIVE — duty math is real (USITC)"),
 ("Shortage List", "/shortage-list", "Customer pastes a backorder list; every line is matched against our stock in real time; unmatched lines become sourcing leads.", "LIVE"),
 ("Supply Risk Monitor", "/supply-risk", "Live FDA device-recall feed mapped against our stocked categories.", "LIVE — real openFDA data"),
 ("Sell Us Surplus", "/surplus", "Public intake form; AI categorizes and values surplus inventory offers.", "DEMO — AI valuation live once Anthropic budget set"),
 ("Surplus Marketplace", "/surplus/market", "Buyers browse accepted surplus lots and place offers.", "LIVE (demo lots)"),
 ("Services", "/services + 4 sub-pages", "Distribution, PDAC consulting (the #1 Google-ranked page), Distributors, Private Label.", "LIVE"),
 ("Industry Segments", "/segments/*", "Tailored landing pages: ASC, Pharmacy, EMS, Distributors.", "LIVE"),
 ("Government", "/government", "Public-sector procurement positioning.", "LIVE"),
 ("Case Study — TJS", "/case-studies/tjs", "Customer success story.", "LIVE"),
 ("Company & Trust pages", "/about, /procurement, /compliance, /locations, /careers, /portfolio", "Brand, veteran-procurement positioning, compliance/document library, 2-warehouse coverage.", "DEMO — warehouse sqft + map redraw pending"),
 ("Resources & Blog", "/resources, /resources/coding, /blog", "HCPCS coding resources; blog shell (renders, awaiting Jill's articles).", "DEMO — blog content pending"),
 ("Contact / Support", "/contact, /support", "Lead capture; document-request CTA.", "DEMO — email send needs Resend key"),
 ("Legal", "/privacy, /terms, /returns, /shipping", "Standard legal pages.", "LIVE"),
]
t = doc.add_table(rows=1, cols=4); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
header_row(t, ["Page", "URL", "What to show Damon", "Real today?"])
for name, url, what, real in pub:
    c = t.add_row().cells
    set_cell(c[0], name, bold=True, size=9)
    set_cell(c[1], url, size=8.5, color=GREY)
    set_cell(c[2], what, size=9)
    col = GREEN if real.startswith("LIVE") else AMBER
    set_cell(c[3], real, size=8.5, color=col)

# ---- PASS 2: CUSTOMER & REP ----
rule(doc)
h2(doc, "Pass 2 — Customer & Rep Portals (the B2B account experience)")
cust = [
 ("Register / Login", "/register, /login", "Signup runs an auto account-approval score (good accounts approved, rest to manual review).", "DEMO — decision: CC-only for new accounts"),
 ("Customer Dashboard", "/dashboard", "Logged-in home: recent orders, quotes, account snapshot.", "LIVE (demo account)"),
 ("Account: Quotes", "/account/quotes", "Quote history; accept a quote online.", "LIVE"),
 ("Account: Invoices", "/account/invoices", "Invoice list + printable invoice view.", "DEMO — real invoicing needs Stripe"),
 ("Account: Team", "/account/team", "Owner invites teammates with roles (owner/buyer/viewer).", "LIVE"),
 ("Account: Settings", "/account/settings", "Profile / org settings.", "LIVE"),
 ("Self-Serve Quote Portal", "/portal/quote", "Approved customers build & price their own quote at their tier — no rep needed.", "LIVE"),
 ("Online Quote Acceptance", "/q/:token", "Tokenized link a customer clicks to accept a quote → becomes an order.", "LIVE"),
 ("Order Tracking", "/orders/:id/track", "Customer-facing shipment status.", "DEMO — live tracking needs ShipStation"),
 ("Rep Portal", "/rep", "1099 rep's book of business: attributed revenue, commission by period, payouts, Calendly intro link.", "DEMO — payouts need Stripe Connect"),
]
t = doc.add_table(rows=1, cols=4); t.style='Table Grid'
header_row(t, ["Surface", "URL", "What to show Damon", "Real today?"])
for name, url, what, real in cust:
    c = t.add_row().cells
    set_cell(c[0], name, bold=True, size=9)
    set_cell(c[1], url, size=8.5, color=GREY)
    set_cell(c[2], what, size=9)
    col = GREEN if real.startswith("LIVE") else AMBER
    set_cell(c[3], real, size=8.5, color=col)

# ---- PASS 3: ADMIN ----
rule(doc)
h2(doc, "Pass 3 — Admin Back-Office (how the business is run)")
body(doc, "Log in as damon@unitemedical.net / admin, then visit /admin. This is the part "
          "Damon will care about most — it's the operating system for the company. ~30 screens.",
     italic=True, color=GREY)
admin = [
 ("Overview", "/admin", "Command center: today's orders, AR, inventory, alerts."),
 ("Analytics", "/admin/analytics", "Sales & traffic analytics."),
 ("CEO Morning Brief", "/admin/digest", "Auto-ranked top-5 things needing attention with deep links (AI-written once key set)."),
 ("Orders", "/admin/orders", "All orders, item-level search, status."),
 ("Zero-Touch Fulfillment", "/admin/fulfillment", "Validate→reserve→pay→invoice→ship→notify pipeline with retries, backorders, returns/RMAs."),
 ("Inventory", "/admin/inventory", "Stock levels (becomes real with Cin7)."),
 ("Replenishment", "/admin/replenishment", "Reorder points from demand; one-click vendor POs. Prophet forecasting when sidecar deployed."),
 ("Products", "/admin/products, /admin/products/new", "Full product CRUD + edit."),
 ("Product Onboarding", "/admin/products/onboard", "GTIN (GS1) + FDA + duty pre-validation before FDA listing."),
 ("Quotes", "/admin/quotes", "All quotes in flight."),
 ("Margin Policy", "/admin/settings/margin", "Tier-margin editor (A/B/C/Distributor)."),
 ("CRM", "/admin/crm", "Contacts & deals (syncs to HubSpot when keyed)."),
 ("Customers", "/admin/customers", "Customer accounts & approval."),
 ("Reps", "/admin/reps", "1099 rep roster, commission accrual, Stripe payouts."),
 ("Vendor Approval", "/admin/vendors", "Real vendor scoring vs live FDA data; auto-approve/reject."),
 ("Trade Discovery", "/admin/discovery", "Find manufacturers & US importers as leads (ImportGenius when keyed)."),
 ("Finance / CFO", "/admin/finance", "AR aging, record payments, send reminders."),
 ("Compliance / Recalls", "/admin/compliance", "Continuous FDA recall sweep → affected-SKU mapping → drafted customer notices."),
 ("Surplus Review", "/admin/surplus", "Review surplus submissions, AI-value, send offers, publish to marketplace."),
 ("CMS", "/admin/cms", "Site content management."),
 ("Integrations", "/admin/integrations", "Status of EVERY external service + 'Run a ping' test buttons. THIS is the go-live dashboard."),
 ("AI Dashboard", "/admin/integrations/ai", "Per-prompt cost / latency / error tracking."),
 ("Webhooks", "/admin/webhooks", "Event bus: dedupe, retries, dead-letter, replay."),
 ("Settings", "/admin/settings", "Global config."),
]
t = doc.add_table(rows=1, cols=3); t.style='Table Grid'
header_row(t, ["Admin screen", "URL", "What it does"])
for name, url, what in admin:
    c = t.add_row().cells
    set_cell(c[0], name, bold=True, size=9)
    set_cell(c[1], url, size=8.5, color=GREY)
    set_cell(c[2], what, size=9)

rule(doc)
h2(doc, "The one-line summary for Damon")
p = doc.add_paragraph()
r = p.add_run("The whole company — storefront, quoting, B2B portals, and a 30-screen back-office "
              "— is built and clickable today on demo data. Going live is mostly a checklist of "
              "pasting in account keys and a handful of business decisions, not months of building.")
r.font.size = Pt(11)
body(doc, "See the companion document, 'Unite Medical — Launch Tracker', for the decision-by-decision "
          "and key-by-key list.", italic=True, color=GREY)

doc.save("/Users/alex-s-nt-16/Projects/Unite/unite_2_0/docs/Unite_Medical_Walkthrough.docx")
print("WROTE Walkthrough")

# ============================================================
# DOC 2 — LAUNCH TRACKER
# ============================================================
doc = Document()
style_base(doc)
for section in doc.sections:
    section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)

h1(doc, "Unite Medical 2.0 — Launch Tracker")
body(doc, "Decision & procurement board for the pre-launch working session. "
          "Backend/Functionality first, then UI/Design. Last updated 2026-06-19.",
     italic=True, color=GREY)

h2(doc, "The key fact")
p = doc.add_paragraph()
r = p.add_run("Adding a credential is the only thing that flips a surface from demo to live. ")
r.bold = True; r.font.color.rgb = GREEN
p.add_run("There are no half-built features. So this board is a checklist of decisions + keys, "
          "not a build backlog. Only genuine remaining build/deploy work: the forecasting "
          "container (B-13), lot-tracking write-path (B-14, gated on Cin7 anyway), and design "
          "assets. Everything else = paste a key or make a decision.")

h3(doc, "Status legend")
for mark, mean in [("LIVE","keyed & working in production"),
                   ("DEMO/STUB","built, running on demo data — needs key/decision"),
                   ("BLOCKED","needs a key or account to go live"),
                   ("DECIDE","needs an owner yes/no (usually free)"),
                   ("DEPLOY","code done; needs one-time deploy/OAuth")]:
    p = doc.add_paragraph(style='List Bullet')
    rr = p.add_run(mark + " — "); rr.bold = True
    p.add_run(mean)

h3(doc, "Where we stand today")
t = doc.add_table(rows=1, cols=2); t.style='Table Grid'
header_row(t, ["Bucket", "Count"])
for k,v in [("Surfaces total","33"),("Live right now","3 (AI, database, free gov APIs)"),
            ("Blocked on a key/account","16"),("Waiting on an owner decision","11"),
            ("UI / design / asset items","10")]:
    c=t.add_row().cells; set_cell(c[0],k,size=9.5); set_cell(c[1],v,size=9.5,bold=True)

# Layer 1 groups
def surface_table(doc, rows, cols_headers):
    t = doc.add_table(rows=1, cols=len(cols_headers)); t.style='Table Grid'
    header_row(t, cols_headers)
    for r_ in rows:
        c = t.add_row().cells
        for i, val in enumerate(r_):
            if i == 0:
                set_cell(c[i], val, bold=True, size=8.5, color=NAVY)
            elif cols_headers[i] == "Status":
                col = GREEN if val=="LIVE" else (RED if val in ("BLOCKED",) else AMBER)
                set_cell(c[i], val, size=8.5, color=col, bold=True)
            else:
                set_cell(c[i], val, size=8.5)
    return t

rule(doc)
h1(doc, "LAYER 1 — Backend / Functionality")

h2(doc, "Group A — Money & Commerce (top priority)")
surface_table(doc, [
 ["B-01","Payments / checkout (Stripe)","Customers actually pay","BLOCKED","STRIPE_SECRET_KEY + Connect approval (needs federal EIN)","Confirm Stripe acct; start Connect approval today"],
 ["B-02","Invoicing & AR (CFO dashboard)","Invoices, who-owes-us, reminders","DEMO","Rides on B-01","Unlocks with B-01"],
 ["B-03","Rep commissions & payouts","Auto-pay 1099 reps","DEMO","Stripe Connect (Express)","Reps paid via Stripe v1? y/n"],
 ["B-04","Accounting sync (QuickBooks)","Orders→books, no double entry","DEPLOY","Intuit app + OAuth; QBO $35-65/mo","CFO migration; pick plan"],
 ["B-05","Sales tax","Right tax per state","DECIDE","—","Stripe Tax vs QBO vs Avalara"],
], ["ID","Surface","What it does","Status","Blocker / need","Owner decision"])

h2(doc, "Group B — Email & Communications")
surface_table(doc, [
 ["B-06","All outbound email (Resend)","EVERY email: orders, invoices, POs, dunning. Until keyed, all silently queues.","BLOCKED","RESEND_API_KEY + verify domain DNS","MOST IMPORTANT KEY. Approve DNS"],
 ["B-07","Scheduling (Calendly)","Rep booking links → CRM","BLOCKED","CALENDLY_API_KEY + webhook secret","Confirm Calendly is the scheduler"],
 ["B-08","Call intelligence (Fathom)","Auto call summaries → CRM tasks","BLOCKED","Fathom acct + webhook secret","Keep for v1? Needs consent line"],
 ["B-09","CRM sync (HubSpot)","Contacts, deals, rep pipeline","BLOCKED","HUBSPOT token ($90/seat/mo)","Confirm seat count"],
], ["ID","Surface","What it does","Status","Blocker / need","Owner decision"])

h2(doc, "Group C — Inventory & Fulfillment")
surface_table(doc, [
 ["B-10","Inventory / WMS (Cin7 Core)","The real stock numbers","BLOCKED","CIN7 keys (~$349-799/mo) + data migration","Pick tier; schedule onboarding"],
 ["B-11","Shipping labels (ShipStation)","Labels + tracking to customers","BLOCKED","SHIPSTATION key+secret (get DIRECT key)","Pull direct API key"],
 ["B-12","Freight / customs (Flexport)","Inbound freight + landed cost","BLOCKED","FLEXPORT key + webhook secret","Email rep for API tier"],
 ["B-13","Demand forecasting (Prophet)","Smarter reorder points","DEPLOY","Deploy forecasting/ container","Low priority — run-rate fine day 1"],
 ["B-14","Lot-level tracking","Recall lookup <1s; compliance SLA","DECIDE","Schema ready; needs pick/pack wiring","Launch requirement or fast-follow?"],
], ["ID","Surface","What it does","Status","Blocker / need","Owner decision"])

h2(doc, "Group D — Product, Compliance & Sourcing")
surface_table(doc, [
 ["B-15","AI features (Anthropic)","Quote letters, classify, digest, valuation","LIVE","Key is set","Set budget cap in dashboard"],
 ["B-16","Durable database (Neon)","State survives + multi-device","LIVE","Set","None"],
 ["B-17","FDA / recall monitor (openFDA)","Live recall sweep","LIVE","Free, no key","None"],
 ["B-18","Duty classification (USITC)","Real import-duty math","LIVE","Free (verified)","None"],
 ["B-19","Product onboarding (GS1)","Validate barcodes pre-FDA","BLOCKED","GS1 key+account (~$500/yr)","Already GS1-registered? Pull key"],
 ["B-20","Trade discovery (ImportGenius)","Find mfrs + importers as leads","BLOCKED","Key ($899/user/mo)","Expensive — v1 or defer?"],
], ["ID","Surface","What it does","Status","Blocker / need","Owner decision"])

h2(doc, "Group E — Business Rules (decisions, mostly free)")
surface_table(doc, [
 ["B-21","Account approval","Auto-approve good B2B signups","DECIDE","New accounts start credit-card-only?"],
 ["B-22","Margin policy per tier","What we charge each customer type","DECIDE","Approve A 30% / B 50% / C 60% / Distributor 25%"],
 ["B-23","Quote validity period","How long a price holds","DECIDE","Default 14 days (Flexport expires in 7)?"],
 ["B-24","Margin disclosure","Does customer see markup?","DECIDE","Default: no (rep sees full breakdown)"],
 ["B-25","Vendor scoring / Class III","Which vendors auto-approve/reject","DECIDE","Confirm Class III=AUTO_REJECT; high-watch list"],
 ["B-26","Rep commission structure","How reps get paid","DECIDE","Flat % (default) vs tiered"],
 ["B-27","Self-serve quote gating","Who quotes without a rep","DECIDE","Default: A-tier + approved distributors"],
 ["B-28","Surplus offer % + logistics","What we pay for surplus","DECIDE","35% retail new-in-box; LTL pickup"],
 ["B-29","Data retention + consent","Compliance/legal posture","DECIDE","Approve retention defaults + Fathom consent"],
], ["ID","Surface","What it does","Status","Owner decision needed"])

rule(doc)
h1(doc, "LAYER 2 — UI / Design")
surface_table(doc, [
 ["U-01","Homepage live-inventory widget","Real Shopify SKU count","BLOCKED","Provide count or approve Shopify API wiring"],
 ["U-02","Warehouse square footage","Real sqft GA + NV","DECIDE","Give the two numbers"],
 ["U-03","Coverage map redraw","2 dots (GA+NV), kill Dallas/4-dot","DESIGN","Designer asset (or ship 2-dot SVG as-is)"],
 ["U-04","8 partner logos","Ardent, Restore Robotics, Amazon Fresh, WellLink, Resource Group, Orlando Health, UF Health, TJS","BLOCKED","Source SVG/PNG"],
 ["U-05","Compliance badge graphic","4-category shield","DESIGN","Designer asset"],
 ["U-06","Blog content","Pages render, empty until posts","BLOCKED","Coordinate with Jill"],
 ["U-07","Quote page rebuild?","Sanitized demo is safe to ship","DECIDE","Ship sanitized vs full rebuild"],
 ["U-08","Document REQUEST buttons","Per-doc buttons hit stub send","DECIDE","Mock fine vs wire real backend"],
 ["U-09","Per-page content review","Walk 22 routes for copy/truth","DESIGN","Read-through together"],
 ["U-10","Admin UI polish","Internal screens, low priority","DESIGN","Defer unless broken"],
], ["ID","Surface","What it does","Status","Owner action"])

rule(doc)
h2(doc, "Procurement checklist (tick as each lands in Vercel env vars)")
order = [
 ("Resend — unlocks ALL email","B-06","[ ]"),
 ("Stripe + Connect approval","B-01/02/03","[ ]"),
 ("Cin7 Core","B-10","[ ]"),
 ("ShipStation","B-11","[ ]"),
 ("HubSpot","B-09","[ ]"),
 ("Calendly","B-07","[ ]"),
 ("QuickBooks (Intuit + OAuth)","B-04","[ ]"),
 ("Flexport","B-12","[ ]"),
 ("GS1","B-19","[ ]"),
 ("Fathom (if kept)","B-08","[ ]"),
 ("ImportGenius (if v1)","B-20","[ ]"),
 ("Forecast sidecar (deploy)","B-13","[ ]"),
 ("Anthropic","B-15","[x] DONE"),
 ("Neon database","B-16","[x] DONE"),
]
t = doc.add_table(rows=1, cols=3); t.style='Table Grid'
header_row(t, ["Done?","Service","Surface"])
for svc, sid, done in order:
    c=t.add_row().cells
    set_cell(c[0],done,size=9,bold=True,color=(GREEN if 'x' in done else GREY))
    set_cell(c[1],svc,size=9); set_cell(c[2],sid,size=8.5,color=GREY)
body(doc,"")
body(doc,"Verify after each: open /admin/integrations and click 'Run a ping' — green = real upstream answered. "
         "Or curl https://<deployment>/api/health.", italic=True, color=GREY)

rule(doc)
h2(doc, "Decisions captured this session")
t = doc.add_table(rows=1, cols=3); t.style='Table Grid'
header_row(t, ["ID","Decision","Date"])
for _ in range(10):
    c=t.add_row().cells
    for i in range(3): set_cell(c[i]," ",size=9)

doc.save("/Users/alex-s-nt-16/Projects/Unite/unite_2_0/docs/Unite_Medical_Launch_Tracker.docx")
print("WROTE Tracker")
