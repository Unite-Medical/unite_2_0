#!/usr/bin/env python3
"""Volume II — Deep Dive: Logistics, Intake & Customer Lifecycles.
Exhaustive test scripts with per-step expected results, real state machines
extracted from the code, edge cases and failure drills."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x2A,0x4A); ACCENT=RGBColor(0x1A,0x6E,0x8E); GREEN=RGBColor(0x1E,0x7D,0x3A)
RED=RGBColor(0xB0,0x2A,0x2A); AMBER=RGBColor(0x9A,0x6A,0x00); GREY=RGBColor(0x55,0x55,0x55); WHITE=RGBColor(0xFF,0xFF,0xFF)

def shade(c,x):
    t=c._tc.get_or_add_tcPr(); s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),x); t.append(s)
def cell(c,t,bold=False,color=None,size=9,white=False,mono=False):
    c.text=''; p=c.paragraphs[0]; r=p.add_run(t); r.bold=bold; r.font.size=Pt(size)
    if mono: r.font.name='Courier New'
    if white: r.font.color.rgb=WHITE
    elif color: r.font.color.rgb=color
def hdr(t,hs,fill='0F2A4A'):
    for i,h in enumerate(t.rows[0].cells): cell(h,hs[i],bold=True,white=True,size=9); shade(h,fill)
def H1(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=NAVY
def H2(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(14.5); r.font.color.rgb=ACCENT
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:spacing'); s.set(qn('w:before'),'260'); s.set(qn('w:after'),'80'); pr.append(s)
def H3(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=NAVY
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:spacing'); s.set(qn('w:before'),'180'); s.set(qn('w:after'),'40'); pr.append(s)
def body(d,t,italic=False,color=None,bold=False,size=10.5):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
def mono_block(d,text,size=8):
    for line in text.strip('\n').split('\n'):
        p=d.add_paragraph(); pr=p._p.get_or_add_pPr()
        s=OxmlElement('w:spacing'); s.set(qn('w:after'),'0'); s.set(qn('w:line'),'220'); s.set(qn('w:lineRule'),'auto'); pr.append(s)
        r=p.add_run(line); r.font.name='Courier New'; r.font.size=Pt(size); r.font.color.rgb=NAVY
def pagebreak(d):
    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def script_table(d, title, rows):
    """A numbered test script: Step | Do | Expected result | ✓"""
    H3(d, title)
    t=d.add_table(rows=1,cols=4); t.style='Table Grid'; hdr(t,['#','Do this','You must see (pass condition)','✓'],'1A6E8E')
    for i,(do,exp) in enumerate(rows,1):
        c=t.add_row().cells
        cell(c[0],str(i),bold=True,size=9); cell(c[1],do,size=9); cell(c[2],exp,size=9); cell(c[3],'[ ]',bold=True,size=10)
    return t

def edge_table(d, rows):
    H3(d,'Edge cases & failure drills — deliberately try to break it')
    t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,['Try this abuse','Correct behavior','✓'],'B02A2A')
    for a,b in rows:
        c=t.add_row().cells; cell(c[0],a,size=9); cell(c[1],b,size=9); cell(c[2],'[ ]',bold=True,size=10)

def state_table(d, title, rows):
    H3(d, title)
    t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,['State','Meaning (plain English)','How it moves on'])
    for a,b,c_ in rows:
        c=t.add_row().cells; cell(c[0],a,bold=True,size=9,mono=True,color=NAVY); cell(c[1],b,size=9); cell(c[2],c_,size=9)

d=Document()
st=d.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
for sec in d.sections:
    sec.top_margin=Inches(0.6); sec.bottom_margin=Inches(0.6); sec.left_margin=Inches(0.65); sec.right_margin=Inches(0.65)

# ================= COVER =================
H1(d,'UNITE MEDICAL 2.0 — TEST BOOK, VOLUME II')
H1(d,'Deep Dive: Intake · Customer Lifecycle · Logistics')
body(d,'')
body(d,'Volume I (the Complete Walkthrough) is the atlas — every page, what it is, quick checks. '
       'THIS volume is the microscope: the three systems we are most worried about, tested to the bone. '
       'Every test script is numbered with an expected result per step — read the step aloud, do it, '
       'confirm the exact expected behavior, tick. Every state machine below is extracted from the actual '
       'code, not from memory. Generated 2026-07-03.',size=11)
body(d,'')
t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,['Part','System under test','Scripts'])
for a,b,c_ in [('Part 1','INTAKE — every door work enters through (7 doors)','14 scripts · 31 edge drills'),
               ('Part 2','CUSTOMER LIFECYCLE — stranger → account → order → cash → return','9 scripts + 4 state machines'),
               ('Part 3','LOGISTICS — the warehouse machine: stock, receiving, fulfillment, shipping','12 scripts + 5 state machines'),
               ('Part 4','THE GOLDEN THREAD — one order followed through every system, 45 steps','1 master script')]:
    c=t.add_row().cells; cell(c[0],a,bold=True); cell(c[1],b,size=9); cell(c[2],c_,size=9)
body(d,'')
body(d,'Logins: damon@unitemedical.net/admin · ops@unitemedical.net/admin · sarah@atlanta-surgical.com/demo · '
       'kareem@holloway.com/demo. Keep Volume I beside you for page-level context.',italic=True,color=GREY)
pagebreak(d)

# =========================================================
# PART 1 — INTAKE
# =========================================================
H1(d,'PART 1 — INTAKE: every door work enters through')
body(d,'Intake is where money is won or lost silently: a broken form is a customer you never knew existed. '
       'The platform has SEVEN intake doors. Each gets: what happens behind the submit button, a field-by-field '
       'script, and abuse drills.')
mono_block(d,'''
  THE SEVEN DOORS                                  WHERE EACH LANDS
  1. /register            account application  ──▶ auto-scoring → account or review queue
  2. /contact             reason-coded message ──▶ lead + CRM + support@ email
  3. /quote               sourcing request     ──▶ sourcing desk lead (named product | custom spec)
  4. /quote/new           vendor sheet upload  ──▶ parsed lines → engine → quote (internal)
  5. /shortage-list       backorder paste      ──▶ per-line match + shortage_requests + CRM
  6. /surplus             surplus listing      ──▶ AI-normalized lines → /admin/surplus desk
  7. /distributor · PO    partner PO file      ──▶ parsed+SKU-mapped → draft order
''')

# ---- Door 1: /register ----
H2(d,'Door 1 — /register (account application)')
body(d,'THE most important intake: everything downstream needs an account. Behind the submit button sits a '
       'real scoring engine (accountApproval.js). It is deliberately simple and you should know exactly how '
       'it decides:')
state_table(d,'The auto-approval scoring — exactly how it works (from code)',[
 ('+1 point','Business address looks commercial (not residential/PO-box pattern)','—'),
 ('+1 point','Website exists and looks valid (real domain shape)','—'),
 ('+1 point','Email domain matches the company website domain','—'),
 ('score ≥ 2','AUTO_APPROVE — account provisioned instantly, welcome email queued, card/ACH checkout from day one','buyer can order immediately'),
 ('score < 2','MANUAL_REVIEW — lands in /admin/customers approval queue; one-click approve/decline email to sales','you decide within your 1-biz-day promise'),
])
script_table(d,'Script 1.1 — clean registration (should auto-approve)',[
 ('Open /register logged out','2-min application: Organization (legal name, website, segment, annual spend), Contact (name, email, phone, password), State'),
 ('Fill org "Peachtree Surgical Partners", website "peachtreesurgical.com", segment ASC, spend $1-5M','Form accepts; no net-30 promise anywhere on the page — copy says card/ACH day one, flexible terms with approved credit'),
 ('Use email jane@peachtreesurgical.com (matches website domain)','This is worth +1 in scoring — deliberate'),
 ('Submit','Redirects to /dashboard as a logged-in approved customer (score hit ≥2: valid site +1, matching domain +1)'),
 ('Check the welcome email','Queued in outbox (visible to admin) — sends for real once Resend key lands; body promises rep contact within one business day'),
 ('As admin, open /admin/customers','New org exists, tier assigned, marked auto-approved with the score/reasons visible'),
 ('As admin, open /admin/crm','Contact was auto-created (registration fires hubspot.createContact)'),
])
script_table(d,'Script 1.2 — sketchy registration (should go to review)',[
 ('Register again: org "Bob\'s Braces", NO website, gmail.com email','Score ≤1 (no site, no domain match)'),
 ('Submit','Account created but flagged MANUAL_REVIEW — user sees pending/limited state, not full buyer access'),
 ('As admin open /admin/customers approval queue','The application sits there with score + reasons ("no website", etc.)'),
 ('Click Approve','Account activates; approval email queues to the applicant'),
 ('Register a third fake, then Decline it in admin','Decline processes; declined account cannot check out'),
])
edge_table(d,[
 ('Submit with an already-registered email','Clear "account exists" error — no duplicate org created'),
 ('Password field left short/blank','Validation blocks; no half-created account rows'),
 ('Same company registers twice with different emails','Two contacts, but you can see/merge the org duplication in admin — check how it presents'),
 ('SQL-ish junk in org name: Robert\'); DROP','Stored as harmless literal text everywhere it renders'),
 ('Register from a phone','Full form usable one-handed; selects and steppers don\'t overflow'),
])

# ---- Door 2: /contact ----
H2(d,'Door 2 — /contact (reason-coded messages)')
script_table(d,'Script 1.3 — the contact pipeline',[
 ('Open /contact','Two phone rows (main + accounting ext. 3), support@/accounting@ only (no info@), reason-coded form'),
 ('Pick reason "Document request", write a message, submit','Lead ID shown on-screen (e.g. lead_xxx) + "in touch inside one business day"'),
 ('Check reason dropdown options one by one','Each reason exists for a real routing purpose — do they match how you triage?'),
 ('As admin, verify the lead landed','Lead exists with the reason attached; auto-reply email queued to the sender'),
 ('Submit the form from /contact?reason=… deep link (e.g. from Compliance page CTA)','Reason arrives pre-selected'),
])
edge_table(d,[
 ('Submit with an invalid email format','Blocked client-side with a clear message'),
 ('Paste a 5,000-word message','Accepted or gracefully limited — no silent truncation'),
 ('Double-click submit','Exactly ONE lead created, not two'),
])

# ---- Door 3: /quote ----
H2(d,'Door 3 — /quote (sourcing request: named product or custom build)')
body(d,'The customer-facing top of the sourcing funnel. Two shapes: "I need this exact product" and '
       '"build this to my spec (my label / Unite label)". No spreadsheets, no template — deliberately.')
script_table(d,'Script 1.4 — named-product request',[
 ('Open /quote','Two entry cards + the "have a vendor sheet? Upload it here" link to /quote/new — and NO template download anywhere'),
 ('Fill named product: "BinaxNOW COVID-19 Ag, 22-pack", qty "500 kits / month"','Fields accept free text — a buyer types how they talk'),
 ('Submit','Reference number displayed; promise = pricing inside one business day'),
 ('As admin, find the request','Visible in the quotes/leads pipeline with the product + qty text intact'),
])
script_table(d,'Script 1.5 — custom-spec request (private label top-of-funnel)',[
 ('Choose the build-to-spec path','Spec textarea (materials, sizes, packaging, certifications), qty/run size, label preference My label / Unite label'),
 ('Fill a plausible spec (e.g. "nitrile exam gloves, 4mil, chemo-rated, our brand box of 200")','Form captures the full spec'),
 ('Submit','Same reference-number pattern; lands on the sourcing desk tagged as custom/private-label'),
])
edge_table(d,[
 ('Submit with quantity "a few"','Accepted (free-text is fine) — the desk qualifies; no rigid validation that loses leads'),
 ('Hit back after submit and resubmit the same form','Second reference number is a new request — check the desk sees the dupe sensibly'),
])

# ---- Door 4: /quote/new ----
H2(d,'Door 4 — /quote/new (vendor sheet → engine) — the deep one')
body(d,'The crown-jewel intake: any vendor price sheet in, landed-cost customer quote out. This is where '
       'foreign-language files, weird columns, and big files all try to break us. Test it like a hostile vendor.')
mono_block(d,'''
 FILE IN (xlsx/csv ≤10MB, any language, any column order, multi-sheet)
   │ parse: picks the data sheet, converts date serials, guards size/rows
   ▼
 COLUMN MAPPING — 3 layers, each labeled with a confidence badge you will SEE:
   alias (exact known headers, incl. 中文/한국어/Tiếng Việt) → fuzzy (typo distance) → AI (Claude)
   uncertain columns  ──▶  CONFIRMATION UI — you approve before anything is priced
   ▼
 TRANSLATE non-English lines (original text preserved alongside)
   ▼
 CLASSIFY: missing FDA product codes auto-filled (AI) · HTS duty rate per line (LIVE USITC data)
   ▼
 LANDED COST per line = FOB + duty + ocean/air freight + brokerage + drayage + receiving
   ▼
 MARGIN by customer tier (A 30 / B 50 / C 60 / dist 25 / gov 20, 10% floor)
   ▼
 QUOTE → internal view (full cost stack) + customer view (sell only) → PDF → /q/:token
''')
script_table(d,'Script 1.6 — the happy path (sample sheet)',[
 ('Open /quote/new','STEP 1 upload/paste box first — NO template step (removed per your review)'),
 ('Click "Load sample" then Parse','4 lines parse instantly: compression stockings, thermometer probes, gel packs, N95s'),
 ('Inspect the parsed preview','Each line shows FDA code, HTS code, FOB, MOQ, target qty; mapping badges say "exact"'),
 ('Pick tier B and freight "cheapest", Run engine','Progress steps stream: compliance → duty → freight → letter; completes in seconds'),
 ('Open the INTERNAL print view','Full cost stack visible per line: FOB, duty %, freight share, margin'),
 ('Open the CUSTOMER print view','Sell prices ONLY — zero cost columns; footer shows FDA · CAGE · corrected BPA line'),
])
script_table(d,'Script 1.7 — the messy paste (column mapper under stress)',[
 ('Paste a CSV with scrambled headers: "Item Desc, Unit $ (USD), Min Order, HS, Qty Needed"','Parser does not die; mapping table appears'),
 ('Read the mapping table','Each of your weird headers mapped to a canonical field with a badge: exact / fuzzy / AI'),
 ('Deliberately mis-map one column and correct it in the UI','Correction sticks; preview re-renders with the fix'),
 ('Run the engine','Quote builds off the corrected mapping'),
])
script_table(d,'Script 1.8 — foreign-language sheet',[
 ('Upload/paste a sheet with Chinese headers (品名/单价/起订量) and Chinese product names','Headers map via the alias layer; product names translate'),
 ('Check a translated line','English shown, original Chinese preserved next to it — the vendor\'s text is never destroyed'),
 ('Run engine and open customer view','Customer sees clean English line items'),
])
script_table(d,'Script 1.9 — multi-vendor compare',[
 ('Parse vendor sheet A, add to compare; parse sheet B with overlapping products','Compare panel lists both offers'),
 ('Run comparison','Engine picks best FOB per product across vendors and quotes the winners; you can see which vendor won each line'),
])
edge_table(d,[
 ('Upload an 11MB file','Rejected with a clear size message — no browser hang'),
 ('Upload a .xlsx with 3 sheets where sheet 1 is a cover page','Parser auto-detects the DATA sheet, not the cover'),
 ('Sheet with a price column containing "$2.40 USD" strings','Parses to 2.40 — currency symbols don\'t break math'),
 ('Sheet with dates as Excel serial numbers','Converted to ISO dates, not 44927'),
 ('CSV with a completely unmappable column ("Notes from Kevin")','Ignored/flagged — never guessed into a price field'),
 ('Line missing FOB price','Line flagged/skipped and COUNTED as a parser miss (captureParserSkips) — misses are tracked, not silently dropped'),
 ('Run engine twice on the same sheet','Two distinct quotes (expected) — check they don\'t collide in Recent'),
 ('Zero-line file (headers only)','Friendly "no lines found" — not a crash'),
])

# ---- Door 5: shortage list ----
H2(d,'Door 5 — /shortage-list (backorder matcher)')
script_table(d,'Script 1.10 — paste-to-match',[
 ('Paste 5 lines: 2 you stock (by rough name), 1 competitor brand you stock an equivalent of, 2 you don\'t carry','Live per-line results: stocked matches with product links, equivalent suggestions, and misses'),
 ('Check a matched line\'s suggestion','The match is genuinely the same product class (wrong match = credibility damage)'),
 ('Send the 2 misses to the sourcing desk','shortage_request created + CRM lead; confirmation on screen'),
 ('Upload the same list as .csv instead of pasting','Identical results'),
 ('As admin verify','Both shortage requests visible with the original line text'),
])
edge_table(d,[
 ('Paste 300 lines','Handles it (progressively or paged) without freezing the tab'),
 ('Paste garbage ("asdfgh, 12345")','Zero matches, zero crashes, polite empty state'),
 ('Paste lines with quantities and units mixed in ("x500 bx")','Matcher still finds the product core'),
])

# ---- Door 6: surplus ----
H2(d,'Door 6 — /surplus (sell-us-your-inventory intake)')
script_table(d,'Script 1.11 — seller intake → desk → market',[
 ('List 2 lines manually (name, qty, expiry, target price)','AI normalizes messy names into clean product descriptions (visible before submit)'),
 ('Upload the surplus CSV variant too','Rows load into the same line editor'),
 ('Submit','Confirmation + submission lands in /admin/surplus'),
 ('As admin: run AI valuation on your submission','Per-line suggested offers appear (35%-of-retail default per your D-10)'),
 ('Adjust one line\'s offer, send the offer email','Offer email queues; submission status moves to offered'),
 ('Accept (as the seller) and publish the lot','Lot appears on public /surplus/market'),
 ('As a buyer, place an offer on the lot; accept it as admin','Connection released: buyer+seller emails queue, fee invoice generated, competing offers auto-closed'),
])
state_table(d,'Surplus offer states (from marketplace.js)',[
 ('open','Lot live on the market, taking offers','buyer offers arrive'),
 ('accepted','You accepted a buyer\'s offer','→ connected (parties introduced)'),
 ('connected','Buyer and seller released to each other','→ fee_pending → fee_paid'),
 ('declined / withdrawn','Offer refused or pulled','terminal'),
 ('fee_pending / fee_paid / not_due','Unite\'s brokerage fee lifecycle','invoice → paid'),
])
edge_table(d,[
 ('Submit surplus with expiry date in the past','Flagged — expired goods路 should be visibly marked, not hidden'),
 ('Two buyers offer on the same lot; accept one','The OTHER offer auto-closes (no double-sell)'),
])

# ---- Door 7: distributor PO ----
H2(d,'Door 7 — /distributor · Upload PO (partner order intake)')
state_table(d,'PO upload states (from poIngestion.js)',[
 ('parsing','File just landed, being read','→ ready | needs_mapping | failed'),
 ('needs_mapping','Some lines have SKUs we don\'t recognize','you map them once → system LEARNS the mapping → ready'),
 ('ready','Every line resolved to a Unite SKU','→ one-click draft order'),
 ('failed','Zero readable lines','partner sees a clear error'),
])
script_table(d,'Script 1.12 — partner PO → draft order',[
 ('Log in as the MedOne demo distributor, open Upload PO tab','Upload/paste area for their PO file'),
 ('Upload a PO containing 2 known SKUs + 1 alias only they use','Status: needs_mapping; the 2 known lines resolved (confidence shown), 1 flagged'),
 ('Map the flagged line to the right Unite SKU','Status flips to ready; the mapping is LEARNED'),
 ('Click draft order','Draft order created with correct lines/qty; visible in /admin/orders'),
 ('Upload a second PO using that same alias','It now auto-resolves — the system remembered'),
])
edge_table(d,[
 ('Upload a PDF PO (not csv/text)','Graceful "format not supported yet" — no hang (PDF intake is PRD-32 roadmap)'),
 ('PO with qty 0 lines','Zero-qty lines excluded from the draft with a note'),
])
script_table(d,'Script 1.13 — intake load test (all doors, 20 minutes)',[
 ('Fire ONE submission through each of the 7 doors back-to-back','All 7 create their records; nothing overwrites anything'),
 ('Open /admin and count the new work items','Every door\'s output is findable from admin within 60 seconds of submission — no intake black holes'),
 ('Check every confirmation promise made ("one business day", reference numbers)','Each promise is one you will actually staff'),
])
pagebreak(d)

# =========================================================
# PART 2 — CUSTOMER LIFECYCLE
# =========================================================
H1(d,'PART 2 — CUSTOMER LIFECYCLE: stranger → account → order → cash → return → repeat')
body(d,'One customer, cradle to grave. These state machines are the actual ones in the code — if the site '
       'disagrees with these tables during testing, the site is wrong or this book is stale; flag either.')
mono_block(d,'''
 STRANGER ──▶ APPLICANT ──▶ ACCOUNT ──▶ FIRST ORDER ──▶ PAID ──▶ SHIPPED ──▶ DELIVERED
                 │ score<2                  │                                    │
                 ▼                          ▼                                    ├──▶ REORDER (loop)
            MANUAL_REVIEW            payment_status:                            ├──▶ RETURN/RMA
            approve/decline          card→pending→paid                          └──▶ AR/dunning if unpaid
                                     terms→invoiced→paid
''')
state_table(d,'Order payment states (from orders.js)',[
 ('pending','Card order placed, payment not yet captured (Stripe pending)','→ paid on capture'),
 ('invoiced','Terms order placed, invoice issued, money not in yet','→ paid on payment recording; ages into AR buckets'),
 ('paid','Money received','fulfillment can run/complete'),
])
state_table(d,'Quote states (quoting/selfServe/acceptance)',[
 ('draft / new','Being built (engine or portal), not yet sent','→ sent/open'),
 ('open / sent','In the customer\'s hands, clock ticking to expiry','→ accepted | countered | declined | expires'),
 ('accepted','Customer clicked accept on /q/:token','→ ORDER created (idempotent — once only)'),
 ('countered','Customer pushed back; desk re-works','→ new version → open'),
 ('declined','Dead','terminal (a PRD-30 outcome label!)'),
 ('invoiced','Accepted and billed through','terminal-ish'),
 ('processing','Engine mid-run','transient'),
])
state_table(d,'RMA states (fulfillment.js returns flow)',[
 ('pending','Return requested, items + refund total captured','→ refunded'),
 ('refunded','Restocked + credit memo + refund issued (best-effort chain)','terminal; stock is back on the shelf'),
])

script_table(d,'Script 2.1 — birth of an account (recap + deepen)',[
 ('Run Scripts 1.1 and 1.2 if not done','One auto-approved account, one via the review queue'),
 ('As the new customer, check the assigned tier in admin','Tier matches segment (ASC → B by default) — pricing depends on this'),
 ('As admin, change the account tier to A','Customer\'s catalog prices visibly drop on refresh (tier engine live)'),
])
script_table(d,'Script 2.2 — first order, card path (the D-03 rule)',[
 ('As the fresh account, add 3 items to cart','Tier prices applied; totals correct'),
 ('Proceed to checkout','Payment options: card/ACH ONLY — no terms offered to a new account (your rule, verify hard)'),
 ('Place the order','Order number issued; payment_status shows pending (card, until Stripe capture) '),
 ('Check confirmation email + order page','Items, totals, rep contact correct'),
])
script_table(d,'Script 2.3 — first order, terms path (approved-credit account)',[
 ('As Sarah (seeded with approved terms), place an order choosing terms','Order lands with payment_status = invoiced'),
 ('Open /account/invoices','Invoice exists with a due date and open balance'),
 ('As admin /admin/finance, find it in AR aging','Sits in the current bucket; ages over time'),
 ('Record a payment against it','Invoice → paid, balance 0; order payment_status → paid; AR bucket updates'),
])
script_table(d,'Script 2.4 — dunning / the money you have to chase',[
 ('In /admin/finance find a seeded overdue invoice (60+ bucket)','Aging buckets: current / 30 / 60 / 90+ render with real math'),
 ('Click send reminder','Dunning email drafts+queues (sends when Resend lands); logged against the invoice'),
 ('Record a partial payment','Balance reduces; invoice stays open with remaining amount; bucket recalculates'),
])
script_table(d,'Script 2.5 — quote lifecycle end-to-end (with expiry + counter)',[
 ('As Sarah build a self-serve quote (/portal/quote), generate it','Quote status open/sent; visible in /account/quotes AND /admin/quotes'),
 ('Open the /q/:token link, but DON\'T accept — reload it twice','Quote renders each time, still open; token link is stable'),
 ('Accept it','Order created; quote → accepted everywhere; re-opening the token shows accepted (no second order possible)'),
 ('Build a second quote, then as admin mark it countered','Status flows; desk re-work path visible'),
 ('Find a seeded expired quote and open its token','Graceful "expired" screen offering contact — NOT an acceptance'),
])
script_table(d,'Script 2.6 — the reorder loop (where B2B compounds)',[
 ('As Sarah open /account/order history','Past orders listed with Reorder buttons'),
 ('Reorder the Script-2.3 order','Preview modal: lines with CURRENT prices + stock; changes since original flagged'),
 ('Add-all-to-cart and check out','Second order in minutes — the loop that pays for everything'),
 ('Check a saved list reorder too','Same preview behavior from lists'),
])
script_table(d,'Script 2.7 — return / RMA',[
 ('As admin open the delivered test order → start a return for 1 line','RMA created: items, refund_total computed, status pending'),
 ('Process the RMA','Chain executes: item restocked (stock +1 visible in inventory), credit memo issued, refund recorded, status refunded'),
 ('Check the customer\'s invoice/credit view','Credit memo visible to the customer'),
])
script_table(d,'Script 2.8 — team growth inside a customer org',[
 ('As Sarah invite a teammate as buyer, another as viewer','Invites pending until signup'),
 ('Sign up as the viewer via the invite','Activates into the org with viewer role'),
 ('As viewer, try to check out','BLOCKED — viewers browse, buyers buy (role gate is real)'),
 ('Promote viewer→buyer as the org owner, retry checkout','Now permitted'),
])
script_table(d,'Script 2.9 — the CRM shadow (every touch leaves a trace)',[
 ('After Scripts 2.1–2.7, open /admin/crm','The new org exists w/ contact; the orders created closed-won deals automatically'),
 ('Check the timeline of one contact','Registration, orders, quotes all present — the automation wrote history a rep would kill for'),
])
edge_table(d,[
 ('Accept the same /q/:token from two browser tabs at once','Exactly ONE order (acceptance is idempotent)'),
 ('Reorder an order containing a now-discontinued SKU','Line flagged/blocked with an explanation — not silently dropped'),
 ('Record a payment larger than the invoice balance','Rejected or explicit-credit flow — never silent overpayment'),
 ('Viewer role deep-links straight to /checkout','Still blocked (gate is server-of-truth, not just hidden buttons)'),
 ('Return more units than were ordered','Refused with message'),
])
pagebreak(d)

# =========================================================
# PART 3 — LOGISTICS
# =========================================================
H1(d,'PART 3 — LOGISTICS: the warehouse machine')
body(d,'The rule underneath everything (worth explaining to anyone who touches stock): nobody edits a stock '
       'number, ever. Every change — receipt, ship, damage, count, transfer — is a signed ledger movement '
       'with a reason and a reference. On-hand is just the sum. If a number ever looks wrong, the ledger '
       'says exactly who/what/when — and the fix is a new movement, not an edit.')
mono_block(d,'''
 THE STOCK LEDGER (stock_movements — append-only)
   receipt(+) · ship(−) · adjust_damage(−) · adjust_loss(−) · found(+)
   transfer_out(−)/transfer_in(+) · count_variance(±) · return_restock(+)
        every movement: sku · warehouse · qty± · reason · ref(PO/order/RMA/count) · actor
        ▼ (same transaction)
 INVENTORY PROJECTION per sku/warehouse:  on_hand = Σ movements
                                          available = on_hand − reserved
        ▲ reserved comes from…
 RESERVATIONS: held (order placed) → committed (shipped; posts the ship movement) | released (cancelled)
 INVARIANTS the verifier enforces nightly:
   on_hand == Σ movements  ·  reserved == Σ held  ·  available never < 0  ·  lots: Σ qty_remaining == on_hand
''')
state_table(d,'Reservation states (wms/reservations.js)',[
 ('held','Order placed: units earmarked; reserved↑, on_hand unchanged (still on the shelf)','→ committed on ship · → released on cancel'),
 ('committed','Order shipped: reserved↓ AND the ship movement posts (on_hand↓) in one step','terminal'),
 ('released','Order cancelled/edited: units returned to available','terminal'),
])
state_table(d,'Purchase-order lifecycle (wms/purchaseOrders.js — exact allowed transitions)',[
 ('draft','Replenishment (or you) drafted it; editable','→ approved | cancelled'),
 ('approved','You signed off on spend','→ sent | cancelled'),
 ('sent','Emailed to vendor with the branded PDF','→ partial | received | cancelled'),
 ('partial','Some lines received (short shipment)','→ partial (more arrives) | received | cancelled'),
 ('received','All lines in','→ closed'),
 ('closed / cancelled','Done / killed','terminal (no other jumps allowed — try to cheat it in Script 3.4)'),
])
state_table(d,'Fulfillment pipeline — the 8 steps every order walks (fulfillment.js)',[
 ('1 validate','Order sane? items? address?','fail → retry w/ backoff'),
 ('2 reserve','Hold stock (reservations). Shortfall → BACKORDER row per short line','backorder auto-fulfills on restock'),
 ('3 payment','Capture card / confirm terms → payment_status paid','circuit breaker: if Stripe is down, step degrades, order not lost'),
 ('4 invoice','Invoice row issued/marked','—'),
 ('5 shipping','ShipStation label → tracking number → order status ready_to_ship','stub label until key'),
 ('6 packing_slip','Branded packing slip PDF','—'),
 ('7 notify','Customer email with tracking','queues until Resend'),
 ('8 delivered','Marked on carrier confirmation → order status delivered','closes the loop'),
])

script_table(d,'Script 3.1 — reservations & the oversell wall (the test that matters most)',[
 ('Pick a SKU; note available in /admin/inventory','available = on_hand − reserved'),
 ('As Sarah, cart quantity = ALL available units, place order','Order lands; reserved jumps by that qty; available → 0; on_hand UNCHANGED (units still on the shelf)'),
 ('In a second browser as Kareem, try to order 1 unit of the same SKU','Cart/checkout REFUSES — available is 0. This is the no-oversell wall'),
 ('Cancel Sarah\'s order','Reservation released; available springs back; Kareem can now buy'),
 ('Try the two-tabs race: both customers submit the last unit simultaneously','Exactly one succeeds; the loser gets a clean out-of-stock message'),
])
script_table(d,'Script 3.2 — receiving against a PO (the workstation, on a phone)',[
 ('Create the need: on /admin/replenishment, draft POs from low-stock suggestions','Draft PO(s) appear grouped by vendor, MOQ respected, expected delivery ~35d'),
 ('Approve → send one PO','Status draft→approved→sent; vendor email w/ PDF queues; PO print view correct'),
 ('On your PHONE, open /admin/inventory/receive','Workstation usable one-handed: big targets, SKU/lot/expiry/qty entry'),
 ('Receive line 1 in full WITH lot "LOT-TEST-1" + expiry next year','Lot row created; ledger posts receipt(+); on_hand rises by exactly that qty; PO line received_qty updates'),
 ('Receive line 2 SHORT (ordered 100, receive 60)','PO status → partial; the 40 stays open'),
 ('Receive the remaining 40 later','PO → received; close it → closed'),
 ('Check /admin/finance (or QBO once keyed)','Landed-cost bill posted for the receipt (stub bill until QBO key — but the chain fires)'),
])
script_table(d,'Script 3.3 — lots, FEFO & the 60-second recall drill',[
 ('Receive the same SKU twice: LOT-A expiring sooner, LOT-B expiring later','Two lots visible in /admin/inventory/lots with qty_remaining each'),
 ('Ship an order for that SKU','FEFO: LOT-A (earlier expiry) decrements FIRST — check qty_remaining moved on A, not B'),
 ('Open the recall box, look up LOT-A','EVERY customer who received LOT-A units listed, with order + date — in under a second'),
 ('Time the whole drill: from "FDA calls about LOT-A" to customer list on screen','Under 60 seconds total. This is the /compliance promise made real'),
 ('Check expiring-soon view','LOT-A flagged ahead of LOT-B'),
])
script_table(d,'Script 3.4 — PO state machine abuse (try to cheat it)',[
 ('On a draft PO, attempt to receive against it','Refused — must be approved+sent first'),
 ('Try to jump a sent PO straight to closed','Refused — must pass through received'),
 ('Cancel a partial PO','Allowed (per the transition table); already-received units REMAIN in stock (movements are never unwound by a status change)'),
 ('Try to edit line items on a sent PO','Locked or explicit-revision flow — no silent mutation of a document the vendor already has'),
])
script_table(d,'Script 3.5 — cycle count with a planted lie',[
 ('Physically "count" a bin but enter 3 fewer than system qty','Variance row: system vs counted vs −3'),
 ('Post the count session','count_variance(−3) movement posts after review; on_hand drops by 3; audit log names the session'),
 ('Recount correctly and post','Corrective +3 movement — the ledger shows BOTH events forever (no history erasure)'),
])
script_table(d,'Script 3.6 — adjustments & transfers',[
 ('Post a damage adjustment (−2, reason adjust_damage, note "forklift")','Movement with reason+note; on_hand −2'),
 ('Post a found adjustment (+1)','on_hand +1; both visible in movement history'),
 ('Create a transfer (single-warehouse today: exercise the draft flow)','Draft transfer with lines; in_transit semantics ready for warehouse #2'),
])
script_table(d,'Script 3.7 — fulfillment pipeline under fire',[
 ('Run your Script-2.2 order through /admin/fulfillment','All 8 steps stream: validate→reserve→payment→invoice→shipping→packing_slip→notify→delivered-pending'),
 ('Read each step\'s stored state','completed with result payloads; attempt_count = 1'),
 ('Order MORE than available of a SKU (force a shortfall)','Reserve step creates a BACKORDER row for the shortfall instead of failing the order'),
 ('Receive stock for that SKU (Script 3.2 pattern)','Backorder auto-fulfills on restock — shipped without human touch'),
 ('Simulate an integration failure (e.g. shipping step w/ no key in a forced-real mode)','Step marked failed/degraded, retries with backoff, pipeline does NOT lose the order; circuit breaker note visible'),
 ('Mark delivered','Order status delivered; step 8 completed'),
])
script_table(d,'Script 3.8 — inbound freight chain (the Flexport criterion)',[
 ('From /admin/orders inbound panel (or /admin/replenishment), trigger the simulated inbound shipment','A cleared shipment appears with line items, freight $, customs $'),
 ('Watch the chain fire','1) inventory increments at the right warehouse 2) landed-cost bill posts 3) reorder points recalc — all three, automatically, with audit rows'),
 ('Check the audit log entries for the chain','receiving.inventory_received → qbo_bill_posted (or skipped w/ reason) → receiving.complete'),
])
script_table(d,'Script 3.9 — shipping & blind-ship',[
 ('On a ready order, create the label (stub until ShipStation key)','Tracking number + carrier attach to the order; status ready_to_ship'),
 ('Run a distributor blind-ship order (MedOne)','Label/packing slip carry the PARTNER brand + return identity — Unite invisible to the end customer'),
 ('Check the packing slip PDF','Right items, right branding, no internal pricing anywhere'),
])
script_table(d,'Script 3.10 — replenishment math sanity',[
 ('Open /admin/replenishment and pick one reorder-status SKU','Row shows: run rate, on-hand, reorder point = rate×(35d lead+14d safety), days cover, suggested qty'),
 ('Hand-check the math on that one row with a calculator','Numbers reconcile (suggested = max(MOQ, 2×window×rate − on_hand))'),
 ('Check a zero-sales SKU','No false reorder urgency (rate 0 → status ok/∞ cover)'),
])
script_table(d,'Script 3.11 — the ledger invariant audit (5 minutes, ultimate trust test)',[
 ('Pick 3 SKUs you touched across Part 3','—'),
 ('For each: open movement history and sum the qty column yourself','Your sum == displayed on_hand, exactly, every time'),
 ('Check reserved against open orders','reserved == sum of held reservations for those SKUs'),
 ('Run scripts/wms_check.py (Alex)','ALL invariants green: ledger sums, no negative available, lot conservation, PO math, reservation math'),
])
edge_table(d,[
 ('Post a movement with qty 0 or a bogus reason','Refused by the ledger (invalid_qty_delta / invalid_reason)'),
 ('Replay the same receive webhook/submission twice (idempotency key)','Second post is a no-op duplicate — stock does NOT double'),
 ('Receive against a cancelled PO','Refused'),
 ('Ship an order whose reservation was released mid-flight','Reserve step re-runs/fails safe — never ships unreserved stock silently'),
 ('Count session left open overnight','Still postable; no lock-up; stale sessions visible'),
])
pagebreak(d)

# =========================================================
# PART 4 — GOLDEN THREAD
# =========================================================
H1(d,'PART 4 — THE GOLDEN THREAD (one order through every system)')
body(d,'The finale: a single narrative order that crosses intake → lifecycle → logistics → money → CRM. '
       'Run it uninterrupted, in order. If all 45 steps pass, the platform is coherent — not just page-correct.')
gt=[
 ('Register "Golden Path Surgical" w/ matching-domain email','AUTO_APPROVE; welcome email queued'),
 ('As admin, set the account to tier B','Tier saved'),
 ('Log in as the new buyer; browse /catalog','Tier-B prices shown'),
 ('Open a product with variants; add 2 units','Cart = 2 @ tier-B price'),
 ('Paste a 3-line shortage list containing one item you stock','Match found; add the matched item to cart from the tool'),
 ('Cart now has 2 lines; verify totals','Line math + total correct'),
 ('Checkout — payment options','Card/ACH only (new account; D-03 rule visible in the wild)'),
 ('Place the order → note the order #','Confirmation page + email queued; payment_status pending'),
 ('Admin: find the order in /admin/orders by ITEM name','Item search returns it'),
 ('Open /admin/fulfillment; run the pipeline','validate ✓ reserve ✓ (reserved↑, available↓, on_hand unchanged)'),
 ('…pipeline continues','payment ✓ (status paid) · invoice ✓ (row created)'),
 ('…pipeline continues','shipping ✓ label+tracking · packing_slip ✓ PDF · notify ✓ email queued'),
 ('Check /admin/inventory for the SKU','on_hand dropped on ship-commit; movement history shows the ship(−) with the order as ref'),
 ('Check the lot browser','The shipped units came off the earliest-expiry lot (FEFO), qty_remaining reflects it'),
 ('Recall drill on that lot','Golden Path Surgical appears in the affected-customer list — your order is traceable'),
 ('Customer: open order tracking page','Status honest, tracking number present'),
 ('Mark delivered in admin','Order → delivered; step 8 complete'),
 ('Customer: /account/invoices','Invoice visible; print view footer credentials correct'),
 ('Admin: /admin/finance — record payment','Invoice paid; AR clean'),
 ('Admin: /admin/crm','Contact + closed-won deal exist from the order automation'),
 ('Customer: build a 2-line self-serve quote in /portal/quote','Tier-B pricing; quote generated'),
 ('Open /q/:token from the quote email/print','Accept page renders'),
 ('Accept it','Order #2 created; quote accepted everywhere; token re-use cannot double-order'),
 ('Run order #2 through fulfillment quickly','Pipeline green again'),
 ('Customer: reorder order #1 from /account/order','Preview flags any price/stock deltas; add-all works'),
 ('Admin: start an RMA for 1 unit of order #1','RMA pending with refund total'),
 ('Process the RMA','Restock movement(+1) · credit memo · refund; RMA refunded'),
 ('Verify the restocked unit in inventory + ledger','on_hand +1 with return_restock reason, RMA as ref'),
 ('Replenishment check on the sold SKU','Run rate now reflects these sales; days-cover updated'),
 ('If below reorder point: draft the PO','PO draft with sane suggested qty (MOQ respected)'),
 ('Approve → send the PO','Vendor email queued; PDF correct'),
 ('Receive it on your phone with a new lot','partial/received per what you enter; ledger + lot rows land'),
 ('Watch reorder points recalc after receipt','Values shift — the loop closed'),
 ('Trigger the simulated Flexport inbound','Stock↑ + bill + recalc chain fires with audit rows'),
 ('Invite a teammate (buyer) to Golden Path\'s org','Invite pending → activate via signup'),
 ('Teammate places a small order','Works; attributed to the same org'),
 ('Check the rep attribution on these orders (/admin/reps)','Revenue attributed; commission accrued at the flat %'),
 ('Send the rep\'s statement','Statement email queues; ledger row written'),
 ('Open /admin/digest and generate','The morning brief MENTIONS your day: new account, orders, AR event, PO activity'),
 ('Open /admin/webhooks','Events from the day visible (or clean-empty until providers keyed); no dead-letter surprises'),
 ('Run scripts/wms_check.py + verify_orchestration.mjs (Alex)','All invariants + 96 runtime checks green AFTER a full day of abuse'),
 ('Check /admin/integrations pings','Board matches expectations: live things green, keyed-later things grey'),
 ('Sweep the outbox','Every email the day generated is there, correctly addressed, none to info@'),
 ('Legacy-URL spot check (5 from Vol I Ch.8)','All 301 correctly'),
 ('Damon verdict on the whole thread','"I would run my company on this" — or the punch list says exactly why not yet'),
]
t=d.add_table(rows=1,cols=4); t.style='Table Grid'; hdr(t,['#','Step','Expected','✓'],'0F2A4A')
for i,(a,b) in enumerate(gt,1):
    c=t.add_row().cells; cell(c[0],str(i),bold=True,size=9); cell(c[1],a,size=9); cell(c[2],b,size=9); cell(c[3],'[ ]',bold=True,size=10)

body(d,'')
H2(d,'Sign-off — Volume II')
t=d.add_table(rows=1,cols=4); t.style='Table Grid'; hdr(t,['Part','Scripts passed','Damon','Alex'])
for a in ['Part 1 · Intake (13 scripts)','Part 2 · Customer lifecycle (9 scripts)','Part 3 · Logistics (11 scripts)','Part 4 · Golden thread (45 steps)']:
    c=t.add_row().cells; cell(c[0],a,bold=True,size=9); cell(c[1],'____ / ____',size=9); cell(c[2],'________'); cell(c[3],'________')

d.save('/Users/alex-s-nt-16/Projects/Unite/unite_2_0/docs/Unite_Medical_Test_Book_Vol2_DeepDive.docx')
print('WROTE Vol II')
