#!/usr/bin/env python3
"""Damon launch handoff — walkthrough-first edition."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x2A,0x4A); ACCENT=RGBColor(0x1A,0x6E,0x8E)
GREEN=RGBColor(0x1E,0x7D,0x3A); RED=RGBColor(0xB0,0x2A,0x2A)
AMBER=RGBColor(0x9A,0x6A,0x00); GREY=RGBColor(0x55,0x55,0x55); WHITE=RGBColor(0xFF,0xFF,0xFF)

def shade(c,x):
    t=c._tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),x); t.append(s)
def cell(c,t,bold=False,color=None,size=9,white=False):
    c.text=''; p=c.paragraphs[0]; r=p.add_run(t); r.bold=bold; r.font.size=Pt(size)
    if white: r.font.color.rgb=WHITE
    elif color: r.font.color.rgb=color
def hdr(t,hs,fill='0F2A4A'):
    for i,h in enumerate(t.rows[0].cells): cell(h,hs[i],bold=True,white=True,size=9); shade(h,fill)
def h1(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
def h2(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(13.5); r.font.color.rgb=ACCENT
    pr=p._p.get_or_add_pPr(); s=OxmlElement('w:spacing'); s.set(qn('w:before'),'220'); s.set(qn('w:after'),'60'); pr.append(s)
def body(d,t,italic=False,color=None,bold=False,size=10.5):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
def bullet(d,t,lead=None):
    p=d.add_paragraph(style='List Bullet')
    if lead: rr=p.add_run(lead); rr.bold=True
    p.add_run(t)

d=Document()
st=d.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
for sec in d.sections:
    sec.top_margin=Inches(0.6); sec.bottom_margin=Inches(0.6)
    sec.left_margin=Inches(0.7); sec.right_margin=Inches(0.7)

h1(d,"Unite Medical 2.0 — Launch Handoff: Your Walkthrough")
body(d,"Damon — your six fixes are applied, verified, and locked in (details at the end). "
       "What stands between us and flipping the domain is one thing: YOU, clicking through YOUR product. "
       "This doc is that walkthrough. Budget 60–90 minutes. Generated 2026-07-03.",italic=True,color=GREY)

h2(d,"Before you start (2 minutes)")
t=d.add_table(rows=1,cols=2); t.style='Table Grid'; hdr(t,["What","Value"])
for k,v in [
 ("Site (staging — the real domain isn't cut over yet)","https://unite-2-0.vercel.app  (Alex pushes the fixes first — confirm with him before starting)"),
 ("Your admin login","damon@unitemedical.net / admin"),
 ("Customer test login","sarah@atlanta-surgical.com / demo"),
 ("How to flag a problem","Text Alex the URL + one line. Don't stop walking — keep a running list."),
 ("What 'demo data' means","Orders, inventory counts, CRM entries are seeded examples until keys/counts land. The PAGES and FLOWS are real."),
]:
    c=t.add_row().cells; cell(c[0],k,bold=True,size=9,color=NAVY); cell(c[1],v,size=9)

h2(d,"Part 1 — Walk it like a stranger (25 min)")
body(d,"Open the site logged OUT. You are a materials director who has never heard of Unite. "
       "At every page ask three questions: Is it true? Would I trust it? Would I buy?")
walk1=[
 ("1","Homepage /","Read the hero out loud. Click every segment card. Does each one land where a buyer expects?"),
 ("2","Catalog → 3 products you know cold","Names, photos, prices, pack sizes. Open 'Stocked equivalents' — are the substitutes sane?"),
 ("3","/services/pdac","Your reputation page — read EVERY word. The #1 Google rank lands people here."),
 ("4","/robotics","The Restore program pitch. You confirmed the claims; now judge the presentation."),
 ("5","/government","Read the new BPA wording you asked for: 'Medava SKUs · via authorized SDVOSB distributor (contract holder).' Exactly right now?"),
 ("6","/about","Your letter (2016 line gone, 500M softened), leadership = you alone now. Final read."),
 ("7","/compliance","Credentials grid + document library. Every REQUEST button promise is yours to keep."),
 ("8","/shortage-list","Paste 5 real product names from memory. Do the matches impress or embarrass?"),
 ("9","/supply-risk","Live FDA recall feed — is this something you'd show a prospect?"),
 ("10","/contact","Submit a real message. Confirm it arrives at support@ (NOT info@ — that's gone)."),
]
t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,["#","Stop","What to do"])
for n,s,w in walk1:
    c=t.add_row().cells; cell(c[0],n,bold=True,size=9); cell(c[1],s,bold=True,size=9,color=NAVY); cell(c[2],w,size=9)

h2(d,"Part 2 — Spend money like a customer (20 min)")
body(d,"This is the pass that matters most. Do it twice — once fresh, once as Sarah (demo customer).")
walk2=[
 ("1","Register a fake company","/register — note the NEW copy: card/ACH from day one, 'flexible terms with approved credit.' No net-30 promise. Feel right?"),
 ("2","Fill a cart","3 items. Watch tier pricing. Try to add something out of stock — it must refuse."),
 ("3","Check out","Complete the order. Confirmation page + email. (Card is test-mode until Stripe key.)"),
 ("4","Track it","Open the order status page a customer would see."),
 ("5","Build a self-serve quote","/portal/quote — add SKUs, generate the quote."),
 ("6","Accept it online","Open the /q/… link like a customer would. Accept. Watch it become an order."),
 ("7","Download the quote PDF","THIS LANDS ON CUSTOMER DESKS. Check the footer: FDA · CAGE · BPA — new wording, no 'SDVOSB partner.'"),
 ("8","Reorder","/account/reorder — the one-click repeat purchase. Smooth?"),
]
t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,["#","Step","What to verify"])
for n,s,w in walk2:
    c=t.add_row().cells; cell(c[0],n,bold=True,size=9); cell(c[1],s,bold=True,size=9,color=NAVY); cell(c[2],w,size=9)

h2(d,"Part 3 — Run the company for 15 minutes (admin)")
body(d,"Log in as yourself. This back office is where you'll live after launch.")
walk3=[
 ("1","/admin","Your command center. Do today's numbers read the way you think?"),
 ("2","/admin/digest","The morning brief — are these 5 bullets what you'd want over coffee?"),
 ("3","/admin/orders","Find YOUR test order from Part 2. Open it."),
 ("4","/admin/fulfillment","Walk that order through the pipeline: reserve → invoice → ship."),
 ("5","/admin/quotes","Find YOUR quote from Part 2."),
 ("6","/admin/finance","AR aging. Record a fake payment against your test order."),
 ("7","/admin/inventory","Demo counts (until opening count) — but click into receive/lots/count screens. Usable on a phone in the warehouse?"),
 ("8","/admin/reps","Commission math on demo data — matches what we agreed?"),
 ("9","/admin/consignment + /admin/udi","The Restore-program screens. Match how the deal actually works?"),
 ("10","/admin/integrations","The go-live board. Green = live, grey = waiting on a key. This is your launch dashboard."),
]
t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,["#","Screen","What to do"])
for n,s,w in walk3:
    c=t.add_row().cells; cell(c[0],n,bold=True,size=9); cell(c[1],s,bold=True,size=9,color=NAVY); cell(c[2],w,size=9)

h2(d,"Your six fixes — applied and locked")
t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,["#","You asked","Status"])
for n,a,s in [
 ("1","BPA → 'Medava SKUs via authorized SDVOSB distributor (contract holder)'","DONE — 5 pages + the PDF footer on every quote/invoice/PO"),
 ("2","Remove info@ everywhere","DONE — support@ / accounting@ only"),
 ("3","Founding year 2019 everywhere","DONE — incl. the hidden '2018' in Google's structured data and 'since 2016' in your letter"),
 ("4","Soften the 500M/largest-drop-shipper claim","DONE — now 'shipped … at national scale', no hard number"),
 ("5","Remove Jackie (card, tagline, any woman-owned)","DONE — all three spots (no woman-owned positioning existed)"),
 ("6","Register = card-first, 'flexible terms with approved credit'","DONE — your exact phrasing; Login page too"),
]:
    c=t.add_row().cells; cell(c[0],n,bold=True,size=9); cell(c[1],a,size=9); cell(c[2],s,size=8.5,color=GREEN,bold=True)
body(d,"Locked: each fix is now a forbidden-pattern rule in the build verifier — if any of these ever "
       "reappears, the site literally cannot build. All checks green: content verifier PASS, lint clean, "
       "96/96 runtime checks, 122 routes prerendered.",italic=True,color=GREY)

h2(d,"After your walkthrough — the go sequence")
t=d.add_table(rows=1,cols=3); t.style='Table Grid'; hdr(t,["Step","Who","What"])
for s,w,x in [
 ("1. Punch-list","You → Alex","Text your flagged items. Alex fixes same-day; anything you accept as-is gets noted."),
 ("2. Money keys","You + Alex","Resend key (all email) + Stripe live key — the two true blockers. $1 test charge, then refund."),
 ("3. 'GO'","You","One word from you."),
 ("4. DNS flip","Alex","unitemedical.net → the new site. Old URLs 301 to protect the PDAC ranking."),
 ("5. First 48h watch","Alex","Search Console + PDAC rank + error logs + first real orders."),
]:
    c=t.add_row().cells; cell(c[0],s,bold=True,size=9,color=NAVY); cell(c[1],w,size=9); cell(c[2],x,size=9)

body(d,"")
body(d,"Damon (CEO) — walkthrough complete, GO given: ______________________   Date: ________",bold=True)
body(d,"Alex (CTO) — punch-list cleared, DNS flipped:  ______________________   Date: ________",bold=True)

d.save("/Users/alex-s-nt-16/Projects/Unite/unite_2_0/docs/Unite_Medical_Launch_Handoff.docx")
print("WROTE Launch Handoff")
