#!/usr/bin/env python3
"""Damon-facing final go-live checklist — v2, facts embedded from site source."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0F,0x2A,0x4A); ACCENT=RGBColor(0x1A,0x6E,0x8E)
GREEN=RGBColor(0x1E,0x7D,0x3A); RED=RGBColor(0xB0,0x2A,0x2A)
AMBER=RGBColor(0x9A,0x6A,0x00); GREY=RGBColor(0x55,0x55,0x55); WHITE=RGBColor(0xFF,0xFF,0xFF)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)
def set_cell(cell,text,bold=False,color=None,size=9,white=False):
    cell.text=''; p=cell.paragraphs[0]; r=p.add_run(text); r.bold=bold; r.font.size=Pt(size)
    if white: r.font.color.rgb=WHITE
    elif color: r.font.color.rgb=color
def header_row(t,headers,fill='0F2A4A'):
    for i,h in enumerate(t.rows[0].cells): set_cell(h,headers[i],bold=True,white=True,size=9); shade(h,fill)
def h1(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
def h2(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(13.5); r.font.color.rgb=ACCENT
    pPr=p._p.get_or_add_pPr(); s=OxmlElement('w:spacing'); s.set(qn('w:before'),'220'); s.set(qn('w:after'),'60'); pPr.append(s)
def h3(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=NAVY
def body(d,t,italic=False,color=None,bold=False,size=10.5):
    p=d.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
def bullet(d,t,lead=None):
    p=d.add_paragraph(style='List Bullet')
    if lead: rr=p.add_run(lead); rr.bold=True
    p.add_run(t)

d=Document()
st=d.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
for sec in d.sections:
    sec.top_margin=Inches(0.6); sec.bottom_margin=Inches(0.6)
    sec.left_margin=Inches(0.7); sec.right_margin=Inches(0.7)

h1(d,"Unite Medical 2.0 — Final Go-Live Checklist (Damon)")
body(d,"Everything below was pulled from the live site source on 2026-07-03 — the numbers, claims, and "
       "promises shown are EXACTLY what the site publishes. You don't have to go find anything: read each "
       "line, and circle TRUE or FIX. Anything marked FIX, Alex corrects before launch. "
       "All automated checks (content verifier, 96 runtime checks, build) already pass — this is the part "
       "only you can sign off.",italic=True,color=GREY)

# ============ PASS 0: FACTS, WITH ACTUAL VALUES ============
h2(d,"Pass 0 — Your credentials, as published (read each value: is it exactly right?)")
body(d,"These appear on /compliance, /government, /procurement, /about, the site header, and on every "
       "quote/invoice/PO PDF that goes to customers.")
t=d.add_table(rows=1,cols=4); t.style='Table Grid'
header_row(t,["Credential","Value AS PUBLISHED","Shown as / caveat","TRUE / FIX"])
for cred,val,sub in [
 ("FDA Establishment Reg.","3015727296","'Device distribution'"),
 ("MSPV BPA","36C24123A0077","'Via authorized SDVOSB partner' — never a self-claim"),
 ("CAGE Code","8MK70","'Federal contracting identifier'"),
 ("DUNS","117553945","'SAM.gov registered'"),
 ("Veteran-Owned","DD214 Verified","'ID.me verified'"),
 ("TAA Compliant","Prioritized","'Country of origin documented'"),
 ("Berry Compliant","Medava PPE","'Buy America Act'"),
 ("PDAC Approved","Credentialed","'All orthotics + RegeniCool™ Pro'"),
]:
    c=t.add_row().cells
    set_cell(c[0],cred,bold=True,size=9,color=NAVY); set_cell(c[1],val,bold=True,size=10)
    set_cell(c[2],sub,size=8.5,color=GREY); set_cell(c[3],"TRUE   /   FIX",size=9)

h3(d,"Company facts, as published")
t=d.add_table(rows=1,cols=3); t.style='Table Grid'
header_row(t,["Fact","AS PUBLISHED","TRUE / FIX"])
for fact,val in [
 ("Address (site footer + all PDFs)","1487 Trae Lane, Lithia Springs, GA 30122"),
 ("Phone (header, contact, PDFs)","833.868.6483 — Accounting & Billing at ext. 3"),
 ("Emails in use","support@ · accounting@ · info@unitemedical.net"),
 ("Founded","2019 (your letter also says 'in supply chain since 2016')"),
 ("Warehouse footprint","ONE warehouse — Lithia Springs, GA, ships all 50 states + territories. (Nevada was REMOVED per PRD-28)"),
 ("Same-day shipping promise","Orders before 2:00 PM EST, Mon–Fri, ship same day"),
 ("Pandemic claim (your letter, /about)","'One of the largest direct-to-patient drop shippers in the country… over 500 million units of tests, PPE, and critical supplies'"),
 ("ISO status (/compliance)","'Pursuing ISO 13485 certification' — not certified"),
 ("Recall/MDR (/compliance)","'Lot-level traceability… we file MDR-eligible reports to the FDA for our own products'"),
]:
    c=t.add_row().cells
    set_cell(c[0],fact,bold=True,size=9,color=NAVY); set_cell(c[1],val,size=9)
    set_cell(c[2],"TRUE / FIX",size=8.5)

h3(d,"People & companies, as published on /about")
t=d.add_table(rows=1,cols=3); t.style='Table Grid'
header_row(t,["Who","Bio AS PUBLISHED","TRUE / FIX"])
for who,bio in [
 ("Damon R. — Founder & CEO","'Veteran. 10+ years in global medical supply chain. Founded and sold an orthotic bracing company. Also operates Unite Pharma and Clyne Health.'"),
 ("Jackie S. — Co-Owner","'Doctor of Podiatric Medicine. Years of clinical and surgical experience across OR and outpatient settings.'"),
 ("Unite Pharma (family co.)","'Multi-state licensed wholesale pharmacy and FDA-registered third-party logistics (3PL) provider'"),
 ("Clyne Health (family co.)","'AI-powered concierge medicine platform'"),
 ("Tagline","'Built by a veteran supply-chain operator and a practicing physician'"),
]:
    c=t.add_row().cells
    set_cell(c[0],who,bold=True,size=9,color=NAVY); set_cell(c[1],bio,size=9)
    set_cell(c[2],"TRUE / FIX",size=8.5)

h3(d,"Robotics program claims (/robotics — these are about Restore, verify with them)")
t=d.add_table(rows=1,cols=2); t.style='Table Grid'
header_row(t,["Claim AS PUBLISHED","TRUE / FIX"])
for claim in [
 "'Restore Robotics holds the industry's ONLY FDA 510(k) clearance for remanufacturing da Vinci Xi and DV5 instruments'",
 "Restore Robotics = manufacturer of record; provides full warranty on remanufactured instruments",
 "Your letter: 'the leadership of Restore Robotics chose us to represent their program'",
]:
    c=t.add_row().cells; set_cell(c[0],claim,size=9); set_cell(c[1],"TRUE / FIX",size=8.5)

h3(d,"⚠ ONE CONTRADICTION WE FOUND — decide now")
t=d.add_table(rows=1,cols=2); t.style='Table Grid'
header_row(t,["The conflict","Your call"],fill='B02A2A')
c=t.add_row().cells
set_cell(c[0],"The /register page promises new accounts 'wholesale pricing, NET-30 TERMS, dedicated rep' — but "
              "we decided (D-03) that new accounts start CREDIT-CARD ONLY, with net-30 via a separate credit "
              "application. One of these has to change: the page copy, or the policy.",size=9)
set_cell(c[1],"[ ] Fix page to CC-first    [ ] Keep net-30 promise (reverse D-03)",size=9,bold=True)

h3(d,"Promises the site makes in YOUR name ('one business day' appears 6 places)")
body(d,"Each is an operational commitment. Tick = you'll staff it.",italic=True,color=GREY)
t=d.add_table(rows=1,cols=3); t.style='Table Grid'
header_row(t,["✓","Promise","Where"])
for prom,where in [
 ("Contact inquiries answered within one business day","/contact"),
 ("New accounts approved within one business day","/register"),
 ("Compliance documents turned around inside one business day","/compliance"),
 ("Quote pricing back inside one business day","/quote (start flow)"),
 ("Robotics team follow-up within one business day","/robotics"),
 ("Same-day pick/pack/ship before 2pm EST","/shipping, /locations, FAQs"),
]:
    c=t.add_row().cells
    set_cell(c[0],"[  ]",bold=True,size=10); set_cell(c[1],prom,size=9); set_cell(c[2],where,size=8.5,color=GREY)

# ============ PASS 1: PAGES ============
def checklist_table(d, rows):
    t=d.add_table(rows=1,cols=4); t.style='Table Grid'
    header_row(t,["✓","Page","Where","What to check"])
    for r_ in rows:
        c=t.add_row().cells
        set_cell(c[0],"[  ]",bold=True,size=10)
        set_cell(c[1],r_[0],bold=True,size=9,color=NAVY)
        set_cell(c[2],r_[1],size=8.5,color=GREY)
        set_cell(c[3],r_[2],size=9)
    return t

h2(d,"Pass 1 — Public pages (read as a customer would)")
checklist_table(d,[
 ("Homepage","/","Hero + stats strip (GA warehouse · Direct manufacturer relationships · 2pm cutoff). Segment cards read right"),
 ("Catalog","/catalog","Spot-check 5 SKUs you know cold: names, images, prices, pack sizes"),
 ("Product page","/products/(any 3)","Descriptions accurate; 'Stocked equivalents' suggestions sensible"),
 ("Quote flow","/quote → /quote/new → /quote/engine","Reads like a customer tool; NO internal jargon (vendor names, margins, tooling)"),
 ("Diagnostics","/diagnostics","New page — the diagnostics line as you want it presented"),
 ("Robotics","/robotics + /services/robotics","THE flagship new page. Every Restore claim verified above; program terms right"),
 ("Shortage List","/shortage-list","Paste a fake backorder list; matches look sane"),
 ("Supply Risk","/supply-risk","Live FDA recall feed loads and is relevant"),
 ("Surplus","/surplus + /surplus/market","Intake form works; offer framing OK; listed lots presentable"),
 ("Services hub","/services + distribution/pdac/distributors/private-label","The services described are the services you sell"),
 ("PDAC page","/services/pdac","YOUR page — read every word. #1 Google rank rides on it"),
 ("Government","/government","BPA/CAGE language exact (per Pass 0); 'Ready to bid RFQs, IFBs, RFPs' — true?"),
 ("TJS case study","/case-studies/tjs","TJS OK being named; story numbers real"),
 ("Segments","/segments/asc ·pharmacy ·ems ·distributors","Each speaks its audience's language"),
 ("About","/about","Founder letter (verified above) + leadership + family companies — final read"),
 ("Procurement","/procurement","Veteran positioning; DD214/ID.me framing comfortable"),
 ("Compliance","/compliance","Credentials grid = Pass 0 table; document library right"),
 ("Locations","/locations","ONE dot (Lithia Springs) + 50-state spokes map — matches reality"),
 ("Resources","/resources + /resources/coding","HCPCS guidance you'd defend"),
 ("Blog","/blog","Empty-but-clean until Jill's posts — acceptable?"),
 ("Contact/Support","/contact · /support","Send a test message; confirm it arrives"),
 ("Careers/Portfolio","/careers · /portfolio","Nothing stale"),
 ("Legal","/privacy /terms /returns /shipping","Skim returns policy — it's the one customers fight over"),
])

h2(d,"Pass 2 — Buy something (the money path)")
body(d,"Once logged out → register fresh; once as demo customer sarah@atlanta-surgical.com / demo.")
checklist_table(d,[
 ("Register","/register","Sign up a fake company; note the net-30 contradiction above — behavior should match your call"),
 ("Cart & pricing","/cart","3 items; tier pricing applied; out-of-stock CANNOT be added"),
 ("Checkout","/checkout","Order end-to-end; confirmation page + email arrive"),
 ("Track & reorder","/orders/(id)/track · /account/reorder","Status presentable; one-click reorder works"),
 ("Self-serve quote","/portal/quote","Build → generate → open accept link (/q/…) → accept → becomes order"),
 ("Quote PDF","/account/quotes","Download the PDF — this lands on customer desks. Footer shows FDA · CAGE · BPA correctly"),
 ("Account pages","/dashboard · /account/*","Orders, invoices, quotes, team invite render"),
 ("Distributor portal","/distributor (order·inventory·po-upload·settlement·shipping·documents)","NEW — the consignment/settlement flow matches how you actually run distributor deals"),
])

h2(d,"Pass 3 — Your side of the house (admin)")
body(d,"Log in as damon@unitemedical.net / admin.")
checklist_table(d,[
 ("Command center","/admin","Today's numbers make sense"),
 ("Morning digest","/admin/digest","The 5 bullets are things you'd want flagged"),
 ("Orders → fulfillment","/admin/orders · /admin/fulfillment","Find your Pass-2 test order; walk it through"),
 ("Inventory & WMS","/admin/inventory (+receive/count/lots/transfers)","Labeled honestly (demo counts until opening count); receiving screen usable on a phone"),
 ("Purchase orders","/admin/purchase-orders","PO board makes sense"),
 ("Finance","/admin/finance","AR aging reads right; record a fake payment"),
 ("Reps & CRM","/admin/reps · /admin/crm","Commission math = what we agreed"),
 ("Consignment & UDI","/admin/consignment · /admin/udi","New robotics-era screens — flows match the Restore program"),
 ("Integrations","/admin/integrations","THE go-live board: green = live, grey = stub"),
])

h2(d,"Pass 4 — Launch blockers (confirm with Alex before DNS flips)")
t=d.add_table(rows=1,cols=3); t.style='Table Grid'
header_row(t,["✓","Item","Why it blocks"])
for item,why in [
 ("RESEND key live + a test email received in your inbox","No email = orders confirm silently to nobody"),
 ("STRIPE live key + one $1 test charge, then refund it","No payments = no revenue"),
 ("Stripe webhook green on /admin/integrations","Orders won't update on payment without it"),
 ("301 spot-checks: /services/dealer, /solutions, /about/veteran-owned, /segments/gov","SEO equity — protects the PDAC ranking"),
 ("Google Search Console verified; sitemap submitted","Watch PDAC rank weekly post-launch"),
 ("Anthropic budget cap set","AI is live; cap prevents surprise bills"),
]:
    c=t.add_row().cells
    set_cell(c[0],"[  ]",bold=True,size=10); set_cell(c[1],item,size=9,bold=True); set_cell(c[2],why,size=9,color=GREY)

h2(d,"Known gaps you WILL see (tracked — don't stop the launch)")
bullet(d,"Partner logo bar: new logos still being sourced.")
bullet(d,"Blog empty until Jill's articles land.")
bullet(d,"Inventory counts are demo numbers until the WMS opening count.")
bullet(d,"Trade discovery (/admin/discovery) on sample data — real feed deferred (D-14).")
bullet(d,"Founder photo placeholder on /about ('Damon on the warehouse floor') — need the real shot.")

h2(d,"Sign-off")
body(d,"Every value above read. Everything marked FIX is listed below / on the back page and either "
       "corrected or explicitly accepted. Ship it.",bold=True)
body(d,"")
body(d,"Damon (CEO): ______________________________   Date: ____________")
body(d,"Alex (CTO):  ______________________________   Date: ____________")

d.save("/Users/alex-s-nt-16/Projects/Unite/unite_2_0/docs/Unite_Medical_GoLive_Checklist.docx")
print("WROTE v2 GoLive Checklist")
